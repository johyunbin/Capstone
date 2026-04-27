"""
인종설 중간보고서 v0 — 학교 양식(중간보고서1) 5장 구조
재작성: 2026-04-27 사람 톤으로 다듬기 + Apple SD Gothic Neo + 페이지 번호.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ─── 폰트: Apple SD Gothic Neo (macOS) + Helvetica Neue / Menlo
KOR = "Apple SD Gothic Neo"
ENG = "Helvetica Neue"
MONO = "Menlo"
FIG_DIR = "/Users/hyunbin/Capstone/experiments/figures"


def set_run_font(run, kor=KOR, eng=ENG, size=11, bold=False, italic=False):
    run.font.name = eng
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), kor)
    rFonts.set(qn("w:ascii"), eng)
    rFonts.set(qn("w:hAnsi"), eng)


def add_run(p, text, **kw):
    r = p.add_run(text)
    set_run_font(r, **kw)
    return r


def add_p(doc, text="", size=11, bold=False, align=None, space_after=4, line_spacing=1.4, indent_first=False, page_break_before=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if page_break_before:
        pf.page_break_before = True
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if indent_first:
        pf.first_line_indent = Cm(0.8)
    if align is not None:
        p.alignment = align
    if text:
        add_run(p, text, size=size, bold=bold)
    return p


def add_mixed_p(doc, runs, align=None, space_after=4, line_spacing=1.5, indent_first=True, page_break_before=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if page_break_before:
        pf.page_break_before = True
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if align is not None:
        p.alignment = align
    if indent_first:
        pf.first_line_indent = Cm(0.8)
    for text, kw in runs:
        add_run(p, text, **kw)
    return p


def add_heading_custom(doc, text, level=1, page_break_before=False):
    sizes = {1: 17, 2: 13, 3: 11.5}
    size = sizes.get(level, 11)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if page_break_before:
        pf.page_break_before = True
        pf.space_before = Pt(0)
    else:
        pf.space_before = Pt(8 if level == 1 else 10)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    add_run(p, text, size=size, bold=True)
    return p


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def fill_cell(cell, text, bold=False, align=None, mono=False, fill=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    eng = MONO if mono else ENG
    kor = MONO if mono else KOR
    add_run(p, text, kor=kor, eng=eng, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shade_cell(cell, fill)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_table_widths(t, widths_cm):
    """python-docx 의 columns[i].width 만으로는 Word 가 자동 layout 으로 무시할 수 있어
    모든 row 의 cell width 까지 명시해야 적용된다."""
    for col, w in zip(t.columns, widths_cm):
        col.width = Cm(w)
    for row in t.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)


def add_figure(doc, path, width_inches=5.5, caption=None):
    if not os.path.exists(path):
        add_p(doc, f"[그림 누락: {path}]", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    # figure 와 caption 을 동일 paragraph 에 넣고 keep_together=True 로 페이지 분리 차단
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(12)
    pf.keep_together = True
    img_run = p.add_run()
    img_run.add_picture(path, width=Inches(width_inches))
    if caption:
        br_run = p.add_run()
        br_run.add_break()
        cap_run = p.add_run(caption)
        set_run_font(cap_run, size=9.5)


def add_page_numbers(doc):
    """푸터에 페이지 번호 가운데 정렬"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_run_font(run, size=10)

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ─────────────────────────────────────────────────
doc = Document()

# ─── 폰트 임베딩 (Windows / Linux 등 폰트 미설치 환경 호환)
def enable_font_embedding(doc):
    settings = doc.settings.element
    for tag in ("embedTrueTypeFonts", "embedSystemFonts", "saveSubsetFonts"):
        existing = settings.find(qn(f"w:{tag}"))
        if existing is None:
            elem = OxmlElement(f"w:{tag}")
            elem.set(qn("w:val"), "true")
            settings.append(elem)


enable_font_embedding(doc)

style = doc.styles["Normal"]
style.font.name = ENG
style.font.size = Pt(11)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rPr.append(rFonts)
rFonts.set(qn("w:eastAsia"), KOR)
rFonts.set(qn("w:ascii"), ENG)
rFonts.set(qn("w:hAnsi"), ENG)

section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.8)

# ═════════════════════════════════════════════════
# 표지
# ═════════════════════════════════════════════════

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "Skew-Aware Stratified Sampling 을 이용한", size=22, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(22)
add_run(p, "벡터 카디널리티 추정 정확도 개선 연구", size=22, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
add_run(p, "인종설 중간보고서", size=16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
add_run(p, "팀명 :  속도는벡터", size=14)

t = doc.add_table(rows=5, cols=2)
t.columns[0].width = Cm(3.5)
t.columns[1].width = Cm(6.0)
fill_cell(t.rows[0].cells[0], "팀     원", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="DDDDDD", size=11)
t.rows[0].cells[1].merge(t.rows[0].cells[0])
fill_cell(t.rows[0].cells[0], "팀     원", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="DDDDDD", size=11)
fill_cell(t.rows[1].cells[0], "팀장", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t.rows[1].cells[1], "박세은", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t.rows[2].cells[0], "팀원", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t.rows[2].cells[1], "강재현", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t.rows[3].cells[0], "팀원", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t.rows[3].cells[1], "이동욱", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t.rows[4].cells[0], "팀원", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t.rows[4].cells[1], "조현빈", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
# 모든 row height 1.0cm + paragraph space — 헤더 ↔ 데이터 row 사이 답답함 해소
from docx.enum.table import WD_ROW_HEIGHT_RULE
for row in t.rows:
    row.height = Cm(1.0)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for cell in row.cells:
        for cp in cell.paragraphs:
            cp.paragraph_format.space_before = Pt(4)
            cp.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

t2 = doc.add_table(rows=2, cols=2)
t2.columns[0].width = Cm(3.5)
t2.columns[1].width = Cm(6.0)
fill_cell(t2.rows[0].cells[0], "지도교수", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t2.rows[0].cells[1], "박광현 교수님", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t2.rows[1].cells[0], "지도 연구원", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t2.rows[1].cells[1], "임채림 석사", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
for row in t2.rows:
    row.height = Cm(1.0)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for cell in row.cells:
        for cp in cell.paragraphs:
            cp.paragraph_format.space_before = Pt(4)
            cp.paragraph_format.space_after = Pt(4)

# 지도교수 표 ↔ "2026 년 4 월" 사이 빈 단락 추가 (사용자: 팀원/지도교수 사이 spacing 과 동일하게)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "2026 년 4 월", size=12)

# ═════════════════════════════════════════════════
# Contents
# ═════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.page_break_before = True
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(36)
add_run(p, "Contents", size=28, bold=True)

contents = [
    ("1.  연구 주제", 1),
    ("2.  연구의 필요성", 1),
    ("3.  연구 내용", 1),
    ("     Ⅰ.  Exqutor 의 두 보완책과 단일 테이블 사각지대", 2),
    ("     Ⅱ.  RQ1 — 사각지대의 구조적 한계 확인", 2),
    ("     Ⅲ.  RQ2 — 두 단계 sanitize 의 직교 측정", 2),
    ("     Ⅳ.  RQ3 — 분포를 모르는 환경으로의 확장 (설계)", 2),
    ("4.  현재 진행 상황", 1),
    ("5.  일정 및 역할 배분", 1),
]
for i, (text, lv) in enumerate(contents):
    p = doc.add_paragraph()
    # Ⅳ. 다음 4. 사이 spacing 을 lv=1 → lv=2 의 spacing 과 일치 (사용자 명시)
    next_lv = contents[i + 1][1] if i + 1 < len(contents) else None
    if lv == 2 and next_lv == 1:
        p.paragraph_format.space_after = Pt(14)
    else:
        p.paragraph_format.space_after = Pt(11 if lv == 1 else 7)
    add_run(p, text, size=14 if lv == 1 else 12, bold=(lv == 1))

# ═════════════════════════════════════════════════
# 1. 연구 주제
# ═════════════════════════════════════════════════
add_heading_custom(doc, "1.  연구 주제", level=1, page_break_before=True)

add_mixed_p(doc, [
    ("이번 연구의 주제는 벡터 데이터베이스에서 거리 기반 조건 ", {}),
    ("WHERE (v <-> q) < D", {"eng": MONO, "kor": MONO}),
    ("의 카디널리티(cardinality) 추정 문제이다. 우리는 최근 발표된 Exqutor (arXiv:2512.09695v2) 시스템이 단일 테이블 vector range query 시나리오에서는 활성화되지 않는다는 사실을 소스 코드를 직접 분석하여 확인하였고, 이 사각지대를 두 단계의 sampling 개선으로 보완할 때 추정 정확도가 어느 정도 개선되는지를 정량적으로 측정하였다. 새로운 알고리즘을 제안하기보다는 기존 시스템의 구조적 한계를 분석하고 단계적인 sanitize 의 효과를 측정하는 데 무게를 두었으며, 따라서 본 보고서의 산출물은 측정 가능한 정확도 개선 수치와 그에 대한 인과적 설명으로 구성된다.", {})
])

add_mixed_p(doc, [
    ("연구는 세 단계의 질문으로 분해된다. ", {}),
    ("RQ1", {"bold": True}),
    (" 은 Exqutor 의 Adaptive Sampling 이 단일 테이블 vector range query 에서 어떤 구조적 한계를 가지는지를 정량 검증한다. ", {}),
    ("RQ2", {"bold": True}),
    (" 는 데이터의 분포를 알 수 있을 때, block sampling bias 제거 (BERNOULLI 교체) 와 data-side k-means stratified sampling 의 두 단계 sanitize 가 카디널리티 추정 정확도를 얼마나 개선하는지를 paired Wilcoxon 검정과 5-seed 95 % 신뢰구간으로 측정한다. ", {}),
    ("RQ3", {"bold": True}),
    (" 는 분포를 사전에 모르는 상황에서 RQ2 의 공간 인식 효과를 어느 정도 회수할 수 있는지를 Recovery Rate 프레임워크로 비교 평가하며, 본 보고서는 7 가지 비교 baseline 의 설계만을 제시한다. 본 중간보고서의 검증 가능 산출물은 RQ1 의 8 지표 negative result, RQ2 의 BERN +9.6 % 와 KM20 +1.64 / +1.76 / +3.07 % 의 외적 타당성, 그리고 Two-Level Decomposition 의 Level 2 단독 +19.6 %p 이며, RQ3 의 측정 결과는 5 월의 비교 실험을 마친 뒤 최종보고서에서 보고한다.", {})
])

# ═════════════════════════════════════════════════
# 2. 연구의 필요성
# ═════════════════════════════════════════════════
add_heading_custom(doc, "2.  연구의 필요성", level=1, page_break_before=True)

add_mixed_p(doc, [
    ("벡터 데이터베이스의 활용 범위는 최근 몇 년 사이 크게 넓어졌다. 대규모 언어 모델의 RAG (Retrieval-Augmented Generation) 파이프라인, 멀티모달 검색, 추천 시스템, 이상 탐지 등 다양한 응용에서 벡터 유사도 검색은 핵심 연산이 되었으며, 단순한 top-k 검색을 넘어 ", {}),
    ("WHERE (v <-> q) < D", {"eng": MONO, "kor": MONO}),
    (" 형태의 거리 임계값 기반 범위 검색과 관계형 연산(조인·집계·필터)이 결합된 분석 쿼리, 즉 VAQ (Vector-Augmented Analytical Query) 형태로 사용되는 경우가 많다.", {})
])

add_mixed_p(doc, [
    ("이러한 쿼리를 효율적으로 실행하려면 옵티마이저가 각 연산의 결과 행 수를 정확하게 추정해야 한다. 카디널리티 추정값은 인덱스 사용 여부, 조인 알고리즘과 순서 등 실행 계획의 거의 모든 결정에 영향을 미치기 때문이다. 그러나 pgvector·VBASE·DuckDB 등 주요 시스템은 카디널리티를 각각 33.3 %, 50 %, 100 % 의 고정 비율로 추정하기 때문에 0.001 % 에서 100 % 까지 변동하는 실제 카디널리티와 큰 괴리가 생기며, 이로 인한 잘못된 plan 선택은 실행 시간을 최대 약 1 만 배 (4 orders of magnitude) 까지 증가시킨다. Exqutor 는 이 한계를 인덱스의 유무에 따라 두 경로로 보완한다. 인덱스가 있을 때는 ECQO (Exact Cardinality Query Optimization) 로 HNSW 인덱스에 range query 를 직접 한 번 실행하여 정확한 카디널리티를 얻고, 인덱스가 없을 때는 Adaptive Sampling 으로 동적 샘플링을 통해 카디널리티를 근사한다. 그러나 Exqutor 의 ", {}),
    ("vector.c", {"eng": MONO, "kor": MONO}),
    (" 를 분석한 결과 Adaptive Sampling 의 핵심 hook 은 다중 테이블 조인 시나리오에서만 활성화되도록 설계되어 있어, 이미지 검색·추천·RAG retrieval 의 근간인 단일 테이블 vector range query 가 보완책에서 배제된다는 사실을 확인하였다. 이는 원논문에 명시되지 않은 사각지대이며, 이번 연구는 그 빈자리에서 단일 테이블에 적합한 sampling 방식을 단계적으로 검증한다.", {})
])

# ═════════════════════════════════════════════════
# 3. 연구 내용
# ═════════════════════════════════════════════════
add_heading_custom(doc, "3.  연구 내용", level=1, page_break_before=True)

# Ⅰ
add_heading_custom(doc, "Ⅰ.  Exqutor 의 두 보완책과 단일 테이블 사각지대", level=2)

add_mixed_p(doc, [
    ("Exqutor 의 두 보완책은 동작 원리 자체가 서로 다르다. ECQO 는 벡터 인덱스가 있을 때 옵티마이저의 plan 수립 단계에서 가벼운 range query 를 실제로 실행하여 정확한 카디널리티를 얻는 방식이다. 인덱스 검색이 매우 빠르다는 점을 활용하므로 오버헤드가 매우 작고, 결과적으로 카디널리티는 추정이 아니라 정확한 값으로 옵티마이저에 전달된다. 반면 Adaptive Sampling 은 인덱스가 없을 때 사용되는 방식으로, ", {}),
    ("TABLESAMPLE", {"eng": MONO, "kor": MONO}),
    (" 절을 통해 일부 행만 추출하여 임계값을 만족하는 비율을 계산하고 그 비율로부터 카디널리티를 근사한다. 두 방법 모두 옵티마이저의 비용 모델에 정확한 정보를 제공함으로써 잘못된 plan 선택을 막는 것이 목적이며, 원논문은 두 방법으로 pgvector 에서 최대 약 1,000 배, VBASE 에서 최대 약 1 만 배의 속도 향상을 보고한다.", {})
])

add_mixed_p(doc, [
    ("그러나 우리가 발견한 설계상의 제약은 Adaptive Sampling 쪽에 집중되어 있다. ", {}),
    ("vector.c", {"eng": MONO, "kor": MONO}),
    (" 의 line 243 에 있는 ", {}),
    ("if (table_count > 2)", {"eng": MONO, "kor": MONO}),
    (" 조건문은 Adaptive Sampling 의 핵심 hook 이 호출되기 위한 진입 조건이고, 따라서 테이블이 두 개 이하인 단일 테이블 쿼리는 이 조건을 통과하지 못한다. 예를 들어 ", {}),
    ("SELECT count(*) FROM partsupp_deep_10_subset_1m WHERE (ps_embedding <-> q) < D", {"eng": MONO, "kor": MONO}),
    (" 와 같이 매우 일반적인 형태의 vector range query 는 hook 경로에서 배제되어 PostgreSQL 의 default selectivity 1/3 로 fall-through 된다. 우리는 이 동작을 직접 소스 검증으로 처음 정량 확인하였으며, Exqutor 가 자신의 multi-table only design intent 로 배제한 사각지대(blind spot)에 해당한다고 본다.", {})
])

add_mixed_p(doc, [
    ("이 사각지대를 다음 세 단계의 연구 질문으로 분석·개선한다. RQ1 에서는 사각지대의 구조적 한계와 설계상의 제약을 직접 소스 검증으로 정량화한다. RQ2 에서는 데이터의 분포를 알 수 있을 때 두 가지 직교적인 sanitize, 즉 block sampling bias 제거 (BERNOULLI 교체) 와 data-side k-means 기반 stratified sampling 의 native 구현이 카디널리티 추정 정확도를 어느 정도 개선하는지를 측정한다. RQ3 에서는 데이터의 분포를 사전에 모르는 상황에서 RQ2 의 공간 인식 효과를 어느 정도 회수할 수 있는지를 Recovery Rate 라는 프레임워크로 평가한다. 본 보고서는 RQ1 과 RQ2 의 결과를 보고하고, RQ3 는 설계 안만을 제시한다.", {})
], page_break_before=True)

add_figure(doc, f"{FIG_DIR}/rq1_motivation/slide6_vector_c_snippet.png", width_inches=6.4,
           caption="그림 1.  vector.c L243 의 hook trigger 사각지대와 본 연구의 두 sanitize 패치 (Pivot A: BERNOULLI 교체, Pivot C: KM20 stratified branch +228 lines)")

# Ⅱ. RQ1
add_heading_custom(doc, "Ⅱ.  RQ1 — 사각지대의 구조적 한계 확인", level=2, page_break_before=True)

add_mixed_p(doc, [
    ("RQ1 의 목표는 Exqutor 의 Adaptive Sampling 이 단일 테이블 vector range query 시나리오에서 실제로 어떤 한계를 가지는지를 직접 소스 검증과 EXPLAIN ANALYZE 로 드러내는 것이다. 검증 과정에서 우리는 다음 네 가지 발견에 도달하였다.", {})
])

add_mixed_p(doc, [
    ("첫째, ", {}),
    ("hook trigger 사각지대", {"bold": True}),
    (" 가 존재한다. ", {}),
    ("vector.c", {"eng": MONO, "kor": MONO}),
    (" line 243 의 ", {}),
    ("if (table_count > 2)", {"eng": MONO, "kor": MONO}),
    (" 조건은 단일 테이블 쿼리의 카디널리티 추정 경로 자체를 배제한다. 우리의 검증 query 는 ", {}),
    ("table_count = 1", {"eng": MONO, "kor": MONO}),
    (" 이므로 hook 이 우회되고 PostgreSQL 의 default selectivity 1/3 으로 fall-through 된다. 이 동작은 Exqutor 원논문의 본문에는 명시되지 않은 설계상의 제약이며, 우리의 직접 소스 검증으로 처음 정량 확인되었다.", {})
])

add_mixed_p(doc, [
    ("둘째, hook 을 강제로 활성화할 경우 ", {}),
    ("plan replacement 부작용", {"bold": True}),
    (" 이 발생한다. 위 hook 을 ", {}),
    ("table_count >= 1", {"eng": MONO, "kor": MONO}),
    (" 로 한 줄 우회하면, 단일 테이블 시나리오에서는 outer query 의 base relation 자체가 ", {}),
    ("Sample Scan", {"eng": MONO, "kor": MONO}),
    (" 으로 plan-tree 가 교체되어 outer query 의 결과가 sample 안의 부분 카운트로 격하된다. 실제 100,000 행이 결과로 나와야 할 ", {}),
    ("SELECT count(*)", {"eng": MONO, "kor": MONO}),
    (" 가 32 를 반환하는 식이다. 다중 테이블 조인에서는 base table 의 sample scan 이 join 결과에 부분적으로만 영향을 주지만, 단일 테이블에서는 outer query 자체의 의미가 깨져 버린다. 따라서 hook 의 단순한 강제 활성화는 해법이 되지 못하며, sampling 방식 자체를 손보아야 한다는 결론으로 자연스럽게 이어진다.", {})
])

add_mixed_p(doc, [
    ("셋째, ", {}),
    ("TABLESAMPLE SYSTEM block bias", {"bold": True}),
    (" 의 문제가 있다. Adaptive Sampling 이 사용하는 ", {}),
    ("TABLESAMPLE SYSTEM", {"eng": MONO, "kor": MONO}),
    (" 은 PostgreSQL 의 8 KB 페이지 블록 단위로 데이터를 추출하기 때문에, 같은 블록 안의 행들은 함께 추출되거나 함께 배제된다. 그 결과 행 단위의 균일성이 깨지고 추정의 분산이 커진다. 이 발견은 RQ2-1 의 sanitize 대상으로 곧장 이어진다.", {})
])

add_mixed_p(doc, [
    ("넷째, ", {}),
    ("query feature 를 사전에 식별하기 어렵다", {"bold": True}),
    (" 는 점이 확인되었다. 4 가지 글로벌 skewness 지표 (Fisher γ, log-Fisher γ, tail ratio P99/P50, Bowley skew) 와 query 별 median Q-error 의 Spearman 상관을 96 차원 DEEP 1M subset 의 6,000 측정에서 검증한 결과, 글로벌 4 지표 × 6 selectivity = 24 조합과 그림 2 의 로컬 4 지표 × 6 selectivity = 24 조합을 합한 총 48 조합 모두에서 절대값 0.2 미만으로 가설이 강하게 기각되었다. 100 query 모두가 Fisher γ < 0 (left-skewed) 이며 tail ratio 도 좁은 범위에 모여 있는 점에서, 고차원 벡터의 거리 집중 현상이 글로벌 지표의 변별력을 평탄화시킨 것으로 해석된다. 이 결과는 query 시점에 분포를 사전 식별하는 단순 휴리스틱이 작동하지 않음을 의미하며, 이번 연구가 RQ2 (distribution-aware) 와 RQ3 (distribution-agnostic) 두 트랙을 모두 다루어야 하는 핵심 근거가 된다. 한편 hook 을 강제 활성화한 직후의 update 경로에서는 q-error 발산, sample_size NaN, SIGSEGV 같은 추가 발산이 관찰되었으며, 본 측정에서는 update 경로를 차단하여 sampling method 자체의 효과만을 분리해서 측정하였다.", {})
], page_break_before=True)

add_figure(doc, f"{FIG_DIR}/rq1_motivation/figure_6_phase5_heatmap.png", width_inches=5.6,
           caption="그림 2.  Phase 5 — 로컬 4 지표 × 6 selectivity = 24 조합 (글로벌 24 조합과 합산하면 총 48 조합) 의 Spearman ρ heatmap (모든 조합에서 |ρ| < 0.2)")

# Ⅲ. RQ2
add_heading_custom(doc, "Ⅲ.  RQ2 — 두 단계 sanitize 의 직교 측정", level=2, page_break_before=True)

add_mixed_p(doc, [
    ("RQ2 는 데이터의 분포를 알 수 있을 때 두 가지 직교적인 sanitize 가 카디널리티 추정 정확도를 얼마나 개선하는지 측정한다. ", {}),
    ("RQ2-1", {"bold": True}),
    (" 은 ", {}),
    ("TABLESAMPLE SYSTEM(p%)", {"eng": MONO, "kor": MONO}),
    (" 를 ", {}),
    ("TABLESAMPLE BERNOULLI(p%)", {"eng": MONO, "kor": MONO}),
    (" 로 교체하여 block bias 를 제거하고 행 단위 균일성을 회복하는 것이며, ", {}),
    ("RQ2-2", {"bold": True}),
    (" 는 k-means (K = 20) 으로 분할한 cluster 별로 균등 샘플을 추출하고 Horvitz-Thompson 가중치를 적용하는 stratified sampling 이다. 두 sanitize 는 직교하므로 ablation matrix 에서 효과를 분리해서 측정할 수 있다.", {})
])

add_heading_custom(doc, "(1)  RQ2-1 검증 — Block bias 제거의 효과", level=3)

add_mixed_p(doc, [
    ("DEEP 1M 데이터셋에서 100 개 쿼리와 6 개 selectivity 구간에 대해 SYSTEM 과 BERNOULLI 만을 변화시킨 paired Wilcoxon 검정을 수행하였다. 표 1 은 EXPLAIN ANALYZE 의 ", {}),
    ("plan_rows", {"eng": MONO, "kor": MONO}),
    (" 를 reference 로 한 Python counterfactual 결과이며, native ", {}),
    ("vector.c", {"eng": MONO, "kor": MONO}),
    (" 한 줄 교체로 측정한 같은 비교에서도 + 3.8 ~ + 9.6 % 의 같은 자릿수 개선이 확인되어 두 측정의 방향이 일치한다.", {})
])

t1 = doc.add_table(rows=7, cols=6)
t1.style = "Light Grid Accent 1"
t1.autofit = False
hdrs = ["Selectivity", "SYSTEM median", "BERNOULLI median", "개선율", "p-value", "승리 (SYS<BERN)"]
for i, h in enumerate(hdrs):
    fill_cell(t1.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="E7EBF2")
rows = [
    ["0.001", "2.5970", "2.5970", "+0.0 %",  "0.596",   "6 / 8"],
    ["0.010", "1.5584", "1.2987", "+20.0 %", "0.240",   "45 / 40"],
    ["0.050", "1.2031", "1.1948", "+0.7 %",  "0.0009",  "58 / 37"],
    ["0.100", "1.2120", "1.1323", "+7.0 %",  "<0.001",  "66 / 31"],
    ["0.300", "1.2095", "1.0794", "+12.0 %", "<0.001",  "76 / 24"],
    ["0.500", "1.1527", "1.0519", "+9.6 %",  "<0.001",  "78 / 22"],
]
for ri, row in enumerate(rows, start=1):
    for ci, val in enumerate(row):
        fill_cell(t1.rows[ri].cells[ci], val, align=WD_ALIGN_PARAGRAPH.CENTER)
set_table_widths(t1, [2.5, 2.7, 2.7, 1.8, 2.0, 4.3])

cap1 = add_p(doc, "표 1. Selectivity 별 SYSTEM 과 BERNOULLI 의 median Q-error 비교\n(DEEP 1M, 100 query, plan_rows 기반 Python counterfactual)",
      bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
cap1.paragraph_format.space_before = Pt(8)

add_p(doc, "")

add_mixed_p(doc, [
    ("Selectivity 0.050 이상의 4 개 구간에서 모두 paired Wilcoxon p < 0.001 로 SYSTEM 의 Q-error 가 BERNOULLI 보다 통계적으로 유의하게 크게 나타났으며, 가장 큰 효과는 s = 0.500 에서 78 / 100 의 SYSTEM 불리, median 기준 9.6 %p 개선이었다. ", {}),
    ("vector.c", {"eng": MONO, "kor": MONO}),
    (" 한 줄 교체만으로 카디널리티 추정 정확도가 측정 가능한 수준으로 개선됨을 보여 준다. 반면 매우 낮은 구간 (s = 0.001, 0.010) 에서는 결과 행 수가 작아 ", {}),
    ("plan_rows", {"eng": MONO, "kor": MONO}),
    (" 의 정수 양자화가 분해능을 제한하기 때문에 통계적 차이가 뚜렷하지 않으며, 이 영역의 효과는 RQ2-2 의 stratified sampling 에서 별도로 다룬다.", {})
], page_break_before=True)

add_figure(doc, f"{FIG_DIR}/rq1_motivation/figure_1_phase4_scatter.png", width_inches=5.5,
           caption="그림 3.  SYSTEM vs BERNOULLI paired Q-error 의 4 selectivity panel (s = 0.05 / 0.10 / 0.30 / 0.50, 2×2 격자)")

add_heading_custom(doc, "(2)  RQ2-2 검증 — Stratified sampling 의 5-seed 평가", level=3, page_break_before=True)

add_mixed_p(doc, [
    ("KM20 stratified sampling 의 BERN 대비 추가 효과를 5 개 seed (setseed 0.1 ~ 0.5) 의 반복 측정으로 검증하였다. 표 2 는 세 데이터셋 (DEEP 1M, DEEP 8M, SIFT 1.5M) 의 selectivity 0.500 / 0.050 / 0.010 구간에 대한 5-seed 평균 개선율과 95 % 신뢰구간이다. 측정은 모두 native ", {}),
    ("vector.c", {"eng": MONO, "kor": MONO}),
    (" 의 STRAT 분기에서 수행되었으며, q-error 산출에는 hook 이 ", {}),
    ("elog", {"eng": MONO, "kor": MONO}),
    (" 로 서버 로그에 기록한 ", {}),
    ("hook_est", {"eng": MONO, "kor": MONO}),
    (" 를 기준으로 사용하였다.", {})
])

t2 = doc.add_table(rows=4, cols=4)
t2.style = "Light Grid Accent 1"
t2.autofit = False
hdrs2 = ["데이터셋", "s = 0.500", "s = 0.050", "s = 0.010"]
for i, h in enumerate(hdrs2):
    fill_cell(t2.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="E7EBF2")
rows2 = [
    ["DEEP 1M",   "+1.64 %  [+1.11, +2.18]", "+1.85 %  (단일 seed)",        "+8.93 %  (단일 seed)"],
    ["DEEP 8M",   "+1.76 %  [+0.65, +2.86]", "+0.55 %  [-3.11, +4.21]",     "-0.71 %  (노이즈 영역)"],
    ["SIFT 1.5M", "+3.07 %  [+2.66, +3.48]", "+4.39 %  [+2.63, +6.15]",     "-0.53 %  (양자화 영역)"],
]
for ri, row in enumerate(rows2, start=1):
    for ci, val in enumerate(row):
        fill_cell(t2.rows[ri].cells[ci], val, align=WD_ALIGN_PARAGRAPH.CENTER)
# 데이터셋(좁게) + 3 selectivity 컬럼(넓게) — 두 줄 wrap 회피
set_table_widths(t2, [2.5, 4.5, 4.5, 4.5])

cap2 = add_p(doc, "표 2. KM20 vs BERN 의 5-seed 평균 개선율과 95 % 신뢰구간\n(native vector.c 측정)",
      bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
cap2.paragraph_format.space_before = Pt(8)

add_p(doc, "")

add_mixed_p(doc, [
    ("DEEP 1M 의 s = 0.500 / 0.300 / 0.100 세 구간은 모두 95 % 신뢰구간의 하한이 양수로 확정되었고, DEEP 8M 의 s = 0.500 신뢰구간 [+0.65, +2.86] 은 1M 의 [+1.11, +2.18] 과 겹쳐 데이터 규모가 달라져도 효과가 유지됨을 보여 준다. SIFT 1.5M 의 s = 0.500 신뢰구간 [+2.66, +3.48] 은 DEEP 의 신뢰구간과 겹치지 않으며, 더 쏠린 분포에서는 KM20 의 효과가 약 두 배로 커지는 것이 관찰된다. 그림 4 의 box plot 은 selectivity 별 BERN 과 KM20 의 분포 차이를 직접 보여 주며, 그림 5 는 세 데이터셋의 외적 타당성 비교 (s = 0.500 의 5-seed 95 % CI) 를, 그림 6 은 selectivity gradient 를 함께 보여 준다.", {})
])

add_figure(doc, f"{FIG_DIR}/rq1_motivation/figure_2_phase6_box.png", width_inches=5.6,
           caption="그림 4.  Phase 6 Step 4 — BERNOULLI vs STRATIFIED (KM20) 100 query × 6 selectivity boxplot (DEEP 1M, native vector.c 측정, y 축 log scale)")

add_figure(doc, f"{FIG_DIR}/rq2_aware/figure_8_cross_dataset_bar.png", width_inches=5.6,
           caption="그림 5.  외적 타당성 — DEEP 1M / DEEP 8M / SIFT 1.5M 의 KM20 s = 0.500 효과 (5-seed 95 % CI)")

add_figure(doc, f"{FIG_DIR}/rq2_aware/figure_7_selectivity_gradient.png", width_inches=5.4,
           caption="그림 6. 세 데이터셋의 selectivity 별 KM20 vs BERN 개선율 (5-seed 95 % CI). X 축은 selectivity 의 로그 스케일 내림차순 (좌측 0.5, 우측 0.01)")

add_heading_custom(doc, "(3)  Two-Level Decomposition — 표본 안정화와 공간 인식의 분리", level=3, page_break_before=True)

add_mixed_p(doc, [
    ("KM20 의 개선분이 어떤 메커니즘에서 비롯되는지를 알아보기 위해 우리는 두 수준으로 분해하였다. ", {}),
    ("Level 1 (표본 안정화)", {"bold": True}),
    (" 은 단순히 데이터를 K = 20 으로 분할하여 비례 배분만 적용해도 발생하는 표본 크기 안정화 효과이고, ", {}),
    ("Level 2 (공간 인식)", {"bold": True}),
    (" 는 cluster 가 거리 기반 (KM20) 일 때만 추가로 발생하는 공간 인식 효과이다. DEEP 1M 의 selectivity 0.010 에서 KM20 과 RANDOM20 의 격차는 19.6 %p 로 측정되었으며, 좁은 selectivity 영역에서 무작위 파티션은 오히려 -10.67 % 로 악화된 반면 KM20 은 +8.93 % 를 유지하였다. 즉 Level 2 가 단독으로 +19.60 %p 를 기여한 셈이다. 이 결과는 우리가 본 연구의 출발에서 세웠던 가설 — ", {}),
    ("데이터의 쏠림이 클수록 공간 인식 sampling 의 가치가 크다", {"italic": True}),
    (" — 의 인과 관계를 직접 보여 준다.", {})
])

add_figure(doc, f"{FIG_DIR}/rq2_aware/figure_9_two_level_decomposition.png", width_inches=5.4,
           caption="그림 7. Two-Level Decomposition — Level 1 (표본 안정화) 과 Level 2 (공간 인식) 의 분리")

add_heading_custom(doc, "(4)  HHI 와 CV — 데이터셋별 효과 격차의 사전 예측", level=3, page_break_before=True)

add_mixed_p(doc, [
    ("SIFT 의 KM20 효과가 DEEP 의 약 두 배에 달하는 이유를 정량적으로 설명하기 위해 우리는 각 데이터셋의 K = 20 cluster 크기 분포를 HHI (Herfindahl-Hirschman Index) 와 CV (변동계수) 로 측정하였다. HHI 는 각 cluster 비율의 제곱합으로, 균일할 때 1/K = 0.05 이며 한 cluster 에 집중될수록 1.0 에 가까워진다. 측정 결과 DEEP 1M 의 cluster 크기 CV 는 0.234 였고 SIFT 1.5M 은 0.394 로 약 68 % 더 쏠려 있었으며, 이러한 쏠림 격차가 KM20 효과의 약 두 배 격차와 함께 커지는 경향을 보였다. 즉 KM20 의 효과는 데이터셋이 가지고 있는 고유한 cluster 쏠림 분포로부터 어느 정도 사전에 예측할 수 있으며, 이는 RQ3 (distribution-agnostic) 단계에서의 비교 baseline 으로도 활용된다.", {})
])

add_figure(doc, f"{FIG_DIR}/rq2_aware/figure_10_cluster_skew.png", width_inches=5.4,
           caption="그림 8. DEEP 1M 과 SIFT 1.5M 의 K = 20 cluster 크기 분포 비교 (HHI / CV 정량)")

# Ⅳ. RQ3
add_heading_custom(doc, "Ⅳ.  RQ3 — 분포를 모르는 환경으로의 확장 (설계)", level=2, page_break_before=True)

add_mixed_p(doc, [
    ("RQ3 는 데이터의 분포를 사전에 모르는 상황에서 사전 학습 없이 RQ2 의 공간 인식 효과를 어느 정도 회수할 수 있는지를 측정한다. 회수율은 ", {}),
    ("Recovery Rate = (방법 X − RANDOM20) / (KM20 − RANDOM20)", {"bold": True}),
    (" 로 정의하고, 다음 세 패러다임의 일곱 가지 방법을 동일한 metric 으로 비교하는 프레임워크를 설계하였다.", {})
])

t_rq3 = doc.add_table(rows=8, cols=3)
t_rq3.style = "Light Grid Accent 1"
t_rq3.autofit = False
hdrs_rq3 = ["패러다임", "이름", "핵심 아이디어"]
for i, h in enumerate(hdrs_rq3):
    fill_cell(t_rq3.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="E7EBF2")
rq3_rows = [
    ["Offline Partition",     "LSH Random Hyperplane",  "랜덤 하이퍼플레인으로 K = 20 분할"],
    ["Offline Partition",     "Random Projection",      "Johnson-Lindenstrauss lemma 기반 저차원 사영"],
    ["Offline Partition",     "Hilbert Curve",          "공간 데이터베이스의 space-filling 을 벡터 영역에 적용"],
    ["Offline Partition",     "Mini-batch K-means",     "1 ~ 5 % 데이터만으로 근사 학습"],
    ["Online Query-Adaptive", "Distance-Shell",         "쿼리 벡터 중심의 동심원 stratification"],
    ["Online Query-Adaptive", "KDE-pilot (Neyman)",     "Pilot sample 로 stratum size 의 최적 배분 추정"],
    ["Weight-based",          "Importance Sampling",    "파티션 없이 가중치만으로 분산 축소"],
]
for ri, row in enumerate(rq3_rows, start=1):
    for ci, val in enumerate(row):
        fill_cell(t_rq3.rows[ri].cells[ci], val, align=WD_ALIGN_PARAGRAPH.CENTER if ci != 2 else WD_ALIGN_PARAGRAPH.LEFT)
# 패러다임(좁게) · 이름(중간) · 핵심 아이디어(넓게) — 사용자 명시
set_table_widths(t_rq3, [4.5, 5.0, 6.5])

cap3 = add_p(doc, "표 3. RQ3 의 7 가지 비교 baseline (3 패러다임)",
      bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
cap3.paragraph_format.space_before = Pt(8)

add_p(doc, "")

add_mixed_p(doc, [
    ("일곱 가지 방법의 설계는 모두 확정된 상태이며, 실제 측정은 5 월의 비교 실험 단계에서 수행한다. 평가는 DEEP 1M / DEEP 8M / SIFT 1.5M 의 세 데이터셋에 대해 5 selectivity × 5 seed 의 동일 조건에서 진행하고, 각 방법의 Recovery Rate 가 0.5 이상이면 분포를 모르는 상황에서도 KM20 의 절반 효과를 회수할 수 있다고 판정한다.", {})
])

# ═════════════════════════════════════════════════
# 4. 현재 진행 상황
# ═════════════════════════════════════════════════
add_heading_custom(doc, "4.  현재 진행 상황", level=1, page_break_before=True)

add_mixed_p(doc, [
    ("4 월 초에는 지도 연구실로부터 Exqutor 의 소스 코드와 실험용 데이터셋, 그리고 실험 서버에 대한 접속 권한을 인계받아 빌드와 환경 세팅을 수행하였다. 4 월 14 일까지 빌드와 권한 설정 그리고 ", {}),
    ("vector.c", {"eng": MONO, "kor": MONO}),
    (" 의 patch 적용을 마쳤고, 같은 시점에 본 연구의 출발 가설이었던 글로벌 skewness 가설을 직접 검증하는 과정에서 가설의 기각과 단일 테이블 사각지대 발견이라는 두 가지 핵심 발견에 도달하였다. 4 월 15 일에는 RANDOM20 대조 실험을 통해 좁은 selectivity 영역에서 19.6 %p 의 격차를 확보하였고, 이로써 본 연구 가설의 인과 관계를 데이터로 뒷받침하였다.", {})
])

add_mixed_p(doc, [
    ("이어서 4 월 16 일에는 DEEP 8M 으로 외적 타당성을 확인하였고, 4 월 17 일에는 SIFT 1.5M 까지 확장하여 5-seed 95 % 신뢰구간을 산출하였다. 4 월 21 일에는 자문위원으로부터 본 연구 방향에 대한 회신을 받았으며, 단일 테이블 시나리오로의 좁힘과 두 sanitize 의 직교 측정 설계가 적절하다는 평가였다. 그 이후로 본 보고서의 작성에 착수하여 4 월 28 일의 중간 발표 및 보고 마감을 준비하고 있다.", {})
])

add_mixed_p(doc, [
    ("실험 측면에서는 본 보고서에 보고된 모든 결과 — 표 1 의 SYSTEM vs BERNOULLI 비교, 표 2 의 5-seed 신뢰구간, 그림 1 ~ 8 의 vector.c 패치 시각화, RQ1 negative result heatmap, scatter / boxplot 직교 측정, 외적 타당성 cross-dataset bar, selectivity gradient, Two-Level decomposition, 그리고 cluster skew 비교 — 가 모두 확보된 상태이다. 모든 측정 데이터와 빌드 산출물은 재현 가능하도록 보존하였으며, 원본 코드와 수정본을 함께 유지하여 추후의 검증과 비교를 용이하게 하였다. RQ3 의 측정 결과는 5 월의 비교 실험을 마친 뒤 최종보고서에서 보고한다.", {})
])

# ═════════════════════════════════════════════════
# 5. 일정 및 역할 배분
# ═════════════════════════════════════════════════
add_heading_custom(doc, "5.  일정 및 역할 배분", level=1, page_break_before=True)

t3 = doc.add_table(rows=8, cols=2)
t3.style = "Light Grid Accent 1"
t3.autofit = False
for i, h in enumerate(["주차 / 기간", "핵심 작업"]):
    fill_cell(t3.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="E7EBF2")
sched = [
    ["W1 ~ W5  (3/2 ~ 4/5)",     "팀 구성, 논문 리딩, 1차 자문, 연구제안서·수행계획서 작성·제출, 연구 방향 확정"],
    ["W6  (4/6 ~ 4/12)",         "실험 서버 권한 인계 및 Exqutor 환경 세팅"],
    ["W7  (4/13 ~ 4/19)",        "vector.c 빌드 완료, RQ1 motivation 발견, RQ2-1·RQ2-2 실험 수행"],
    ["W8  (4/20 ~ 4/26)",        "외적 타당성 실험 (DEEP 8M, SIFT 1.5M), 2차 자문 회신 반영"],
    ["W9  (4/27 ~ 5/3)",         "중간보고서 및 중간발표 자료 작성·제출"],
    ["W10 ~ W12  (5/4 ~ 5/24)",  "RQ3 일곱 가지 sampling 방법 비교 실험 수행 및 결과 분석"],
    ["W13 ~ W15  (5/25 ~ 6/14)", "최종발표 (5/27 ~ 5/29), 전시회 (6/5), 최종보고서 작성·제출 (6/11)"],
]
for ri, row in enumerate(sched, start=1):
    for ci, val in enumerate(row):
        fill_cell(t3.rows[ri].cells[ci], val, align=WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT)
set_table_widths(t3, [4.8, 11.2])

cap4 = add_p(doc, "표 4. 캡스톤 2026-1 학기 전체 일정",
      bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
cap4.paragraph_format.space_before = Pt(8)

add_p(doc, "")

t4 = doc.add_table(rows=5, cols=2)
t4.style = "Light Grid Accent 1"
t4.autofit = False
for i, h in enumerate(["팀원", "주요 역할"]):
    fill_cell(t4.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="E7EBF2")
roles = [
    ["박세은 (팀장)",       "전체 일정 관리, 자문 회신 정리, 4 인 합의 주재"],
    ["강재현 (주 발표자)",   "중간발표 슬라이드 검수, 리허설 진행, Q&A 사회"],
    ["조현빈 (실험·문서화)", "RQ1/RQ2 실험 구현, vector.c native 구현 (228 줄), 보고서 작성"],
    ["이동욱 (분석·작도)",   "통계 분석 (CI·Two-Level), 그림 작성, RQ3 설계"],
]
for ri, row in enumerate(roles, start=1):
    for ci, val in enumerate(row):
        fill_cell(t4.rows[ri].cells[ci], val, align=WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT)
set_table_widths(t4, [3.5, 12.5])

cap5 = add_p(doc, "표 5. 팀원별 역할 배분",
      bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
cap5.paragraph_format.space_before = Pt(8)

add_p(doc, "")

add_mixed_p(doc, [
    ("중간발표 이후의 5 월은 RQ3 의 일곱 가지 sampling 방법을 동일한 측정 파이프라인에서 비교하는 작업에 집중하며, 6 월에는 최종발표·전시회·최종보고서 작성을 마무리한다.", {})
])

# ═════════════════════════════════════════════════
add_page_numbers(doc)

from datetime import datetime
ts = datetime.now().strftime("%Y%m%d_%H%M%S")
out = "/Users/hyunbin/Capstone/submission/_drafts/속도는벡터_중간보고서_v1.docx"
doc.save(out)
print(f"saved: {out}")
