"""
인종설 중간보고서 v2 docx 빌더
원본: 인종설 중간보고서.docx (이동욱, 2026-04-27 11:33)
수정: 조현빈 평가 7건 반영 (records/kakaotalk/20260427_중간보고서 docx 평가.md)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

KOR_FONT = "맑은 고딕"
ENG_FONT = "Calibri"


def _apply_font(run, bold=False, size=11, italic=False, mono=False):
    run.font.name = ENG_FONT if not mono else "Consolas"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    target = "Consolas" if mono else KOR_FONT
    rFonts.set(qn("w:eastAsia"), target)
    rFonts.set(qn("w:ascii"), "Consolas" if mono else ENG_FONT)
    rFonts.set(qn("w:hAnsi"), "Consolas" if mono else ENG_FONT)


def add_p(doc, runs, style=None, align=None):
    """runs: list of (text, kwargs)"""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align:
        p.alignment = align
    for text, kw in runs:
        r = p.add_run(text)
        _apply_font(r, **kw)
    return p


def add_h(doc, text, level=1):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    _apply_font(r, bold=True, size=14 - level)
    return h


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def fill_cell(cell, text, bold=False, align=None, mono=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    r = p.add_run(text)
    _apply_font(r, bold=bold, size=10, mono=mono)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shade_cell(cell, fill)


# ─────────────────────────────────────────────────
doc = Document()

# 본문 기본 스타일
style = doc.styles["Normal"]
style.font.name = ENG_FONT
style.font.size = Pt(11)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn("w:rFonts")) or OxmlElement("w:rFonts")
if rFonts.getparent() is None:
    rPr.append(rFonts)
rFonts.set(qn("w:eastAsia"), KOR_FONT)
rFonts.set(qn("w:ascii"), ENG_FONT)
rFonts.set(qn("w:hAnsi"), ENG_FONT)

# 페이지 여백
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ─────────────────────────────────────────────────
# 도입부
# ─────────────────────────────────────────────────

# ¶1 — 카디널리티 추정의 중요성
add_p(doc, [
    ("벡터 데이터베이스에서 ", {}),
    ("WHERE (v <-> q) < D", {"mono": True}),
    (" 와 같은 유사도 임계값 기반의 거리 조건을 처리할 때, 조건에 부합하는 결과 행의 수를 의미하는 카디널리티(cardinality)를 추정하는 것은 매우 중요하다. 옵티마이저는 이 추정치를 바탕으로 인덱스 사용 여부, 다중 테이블 조인 방식, 스캔 전략 등을 결정하므로, 카디널리티의 정확도는 벡터 데이터베이스의 질의 성능에 직접적인 영향을 준다.", {}),
])

# ¶2 — 고정 비율 추정의 한계 (수치 명시 + 4 orders of magnitude)
add_p(doc, [
    ("기존의 상용 벡터 데이터베이스 가운데 pgvector 는 카디널리티를 항상 전체 행 수의 33.3 % 로 고정 추정하고, VBASE 는 50.0 %, DuckDB 는 100 % (즉 필터 효과가 없다고 가정) 로 추정한다. 그러나 실제로 벡터 유사도 검색의 카디널리티는 임계값과 데이터 분포에 따라 0.001 % 에서 100 % 까지 극단적으로 넓은 범위를 가지므로, 이러한 고정 비율 추정은 실제 데이터 크기 대비 수십~수천 배의 오차를 발생시킨다. 그 결과 옵티마이저는 잘못된 실행 계획(예: 불필요한 Full Table Scan 과 Hash Join)을 선택하게 되며, 이로 인한 질의 실행 시간 증가는 최대 4 orders of magnitude (약 10,000 배) 에 달한다.", {}),
])

# ¶3 — Exqutor 의 두 경로 [수정 1]
add_p(doc, [
    ("최근 발표된 Exqutor 시스템은 이 한계를 인덱스의 유무에 따라 두 경로로 해결한다. 인덱스가 있을 때는 ", {}),
    ("ECQO(Exact Cardinality Query Optimization)", {"bold": True}),
    (" 으로 HNSW 인덱스에 range query 를 실제로 한 번 실행하여 ", {}),
    ("정확한 카디널리티", {"bold": True}),
    (" 를 얻고(오버헤드 약 1.88 ms), 인덱스가 없을 때는 ", {}),
    ("Adaptive Sampling", {"bold": True}),
    (" 으로 모멘텀 기반의 동적 샘플링을 통해 카디널리티를 ", {}),
    ("근사", {"bold": True}),
    ("한다. 즉 두 방법은 정확 계산(ECQO)과 근사 추정(Adaptive Sampling)이라는 서로 다른 동작 원리를 가진다. 그러나 Exqutor 의 두 방법, 특히 Adaptive Sampling 에는 다음과 같은 구조적 한계와 사각지대(blind spot)가 존재한다.", {}),
])

# ¶4 — 두 메커니즘 분리 [수정 2]
add_p(doc, [
    ("첫째, Adaptive Sampling 의 정확도 저하는 두 별도 메커니즘에서 발생한다. ", {}),
    ("(a) Block sampling bias", {"bold": True}),
    (" — 기본 샘플링 연산자 ", {}),
    ("TABLESAMPLE SYSTEM", {"mono": True}),
    (" 은 PostgreSQL 의 8 KB 페이지 블록 단위로 데이터를 추출하기 때문에, 같은 블록에 속한 행들이 함께 추출되거나 함께 배제되어 행 단위 균일성이 깨진다. ", {}),
    ("(b) Skewed cluster underrepresentation", {"bold": True}),
    (" — 표본 내 행 분포가 무작위 균일이라 하더라도, 데이터 자체가 특정 영역에 모여 있는(skewed) 경우 무작위 샘플은 작은 cluster 의 비례 배분을 대표하지 못하여 small cluster underrepresentation 이 발생한다. 두 메커니즘은 본 연구의 RQ2-1(BERN sanitize)과 RQ2-2(stratified sampling)로 각각 직교 측정된다.", {}),
])

# ¶5 — Hook 사각지대 + design constraint [수정 3]
add_p(doc, [
    ("둘째, Adaptive Sampling 의 핵심 hook 인 ", {}),
    ("pgvector_set_baserel_rows_estimate_hook", {"mono": True}),
    (" 은 ", {}),
    ("vector.c", {"mono": True}),
    (" 파일 line 243 의 ", {}),
    ("if (table_count > 2)", {"mono": True}),
    (" 조건을 통과해야만 활성화된다. 즉 테이블 개수가 2 개 이하인 단일 테이블 쿼리(예: ", {}),
    ("SELECT count(*) FROM partsupp_deep_10_subset_1m WHERE (ps_embedding <-> q) < D", {"mono": True}),
    (")는 hook 경로에서 배제되고, PostgreSQL 의 default selectivity 인 1/3 로 fall-through 된다. 이 조건을 단일 테이블에서 강제로 활성화하기 위해 ", {}),
    ("table_count >= 1", {"mono": True}),
    (" 로 한 줄 수정하면, base relation 노드 자체가 ", {}),
    ("Sample Scan", {"mono": True}),
    (" 으로 plan-tree 가 교체되어 outer query 결과가 sample 안 부분 카운트로 격하된다(예: 실제 100,000 인 query 에서 ", {}),
    ("SELECT count(*)", {"mono": True}),
    (" 가 32 를 반환). 이외에도 q-error inf 발산, sample_size NaN 발산, Adaptive update path 의 SIGSEGV 등 ", {}),
    ("다섯 가지 design constraint", {"bold": True}),
    (" 가 발견되었으며, 이는 Exqutor 의 Adaptive Sampling update path 가 multi-table only 라는 design intent 위에 만들어진 결과로 해석된다.", {}),
])

# ¶6 — RAG 사각지대
add_p(doc, [
    ("특히 이미지 검색, 추천 시스템, RAG retrieval 과 같이 실무적으로 중요한 기능의 근간은 단일 테이블 vector range query 이며, 이 시나리오가 Exqutor 의 hook 경로에서 배제된다는 점은 학술적·실무적으로 중요한 사각지대(blind spot)이다.", {}),
])

# ¶7 — RQ 정의
add_p(doc, [
    ("본 연구는 다음 세 연구 질문(Research Question)을 통해 Exqutor 의 단일 테이블 사각지대와 분포 처리 한계를 차례로 분석·개선하고자 한다.", {}),
])

add_p(doc, [
    ("RQ1.", {"bold": True}),
    (" Exqutor 의 Adaptive Sampling 이 단일 테이블 vector range query 시나리오에서 어떤 구조적 한계와 design constraint 를 가지는지 직접 소스 검증과 EXPLAIN ANALYZE 로 정량 드러낸다.", {}),
], style="List Bullet")

add_p(doc, [
    ("RQ2.", {"bold": True}),
    (" 단일 테이블에서 데이터 분포를 알 때, block sampling bias 제거(BERNOULLI 교체)와 data-side k-means K=20 기반 stratified sampling 의 두 sanitize 가 카디널리티 추정 정확도를 얼마나 개선하는지 직교 ablation 으로 측정한다.", {}),
], style="List Bullet")

add_p(doc, [
    ("RQ3.", {"bold": True}),
    (" 데이터 분포를 사전에 모르는 상황에서, 사전 학습 없이 RQ2 의 공간 인식 효과를 얼마나 회수할 수 있는지 Recovery Rate 프레임워크의 7-way 비교로 검토한다(본 중간보고서에서는 설계만 보고하며, 실험은 최종보고서 단계에서 수행).", {}),
], style="List Bullet")

# ─────────────────────────────────────────────────
# 연구 및 실험 방법
# ─────────────────────────────────────────────────
add_h(doc, "연구 및 실험 방법", level=1)

# RQ1 검증 [수정 4]
add_h(doc, "RQ1 검증 — 단일 테이블 사각지대의 네 가지 design constraint", level=2)

add_p(doc, [
    ("본 RQ 는 Exqutor 의 Adaptive Sampling 이 단일 테이블 vector range query 시나리오에서 도달조차 못하는 사각지대와, hook 을 강제 활성화했을 때 발생하는 부작용·발산을 직접 소스 검증으로 정량화한다. 본 팀은 다음 네 가지 design constraint 를 발견하였다.", {}),
])

add_p(doc, [
    ("(i) Hook trigger 사각지대.", {"bold": True}),
    (" 앞서 도입부에서 기술한 바와 같이, ", {}),
    ("vector.c", {"mono": True}),
    (" line 243 의 ", {}),
    ("if (table_count > 2)", {"mono": True}),
    (" 조건은 단일 테이블 쿼리의 카디널리티 추정 경로 자체를 배제한다. 본 연구의 검증 query 는 ", {}),
    ("table_count = 1", {"mono": True}),
    (" 이므로 hook 이 우회되어 PostgreSQL default selectivity 1/3 로 fall-through 된다. 이는 Exqutor 원논문 본문에 명시되지 않은 design constraint 이며, 본 팀의 직접 소스 검증으로 처음 정량 확인되었다.", {}),
])

add_p(doc, [
    ("(ii) Plan replacement 부작용.", {"bold": True}),
    (" 위 hook 을 ", {}),
    ("table_count >= 1", {"mono": True}),
    (" 로 한 줄 우회하면, 단일 테이블 시나리오에서는 outer query 의 base relation 자체가 ", {}),
    ("Sample Scan", {"mono": True}),
    (" 으로 plan-tree 가 교체되어(", {}),
    ("Sampling Method: system, Sampling Parameters: ['0.0385'::real]", {"mono": True}),
    (") outer query 결과가 sample 안 부분 카운트로 격하된다. multi-table join 에서는 base table sample scan 이 join 결과의 부분 영향만 주지만, 단일 테이블에서는 outer query 자체의 의미가 깨진다.", {}),
])

add_p(doc, [
    ("(iii) TABLESAMPLE SYSTEM block bias.", {"bold": True}),
    (" Adaptive Sampling 이 사용하는 ", {}),
    ("TABLESAMPLE SYSTEM", {"mono": True}),
    (" 은 PostgreSQL 의 8 KB 페이지 블록 단위로 데이터를 추출한다. 같은 블록 안의 행들은 함께 추출되거나 함께 배제되므로 행 단위 균일성이 깨지며, 이는 RQ2-1 의 sanitize 대상이 된다.", {}),
])

add_p(doc, [
    ("(iv) Query feature 사전 식별 불가능성.", {"bold": True}),
    (" 본 팀은 글로벌 skewness 지표(Fisher γ, log-Fisher γ, tail ratio P99/P50, Bowley skew)와 query 별 median Q-error 사이의 Spearman 상관을 6,000 측정(96 차원 DEEP 1M subset, 100 query × 6 selectivity × 5 seed × 2 mode)에서 검증한 결과, 모든 24 개 조합(4 지표 × 6 선택도)에서 절대값 0.2 미만으로 유의미한 상관이 없음을 확인하였다. 이는 query 시점에 분포 특징을 사전 식별하는 단순 휴리스틱이 작동하지 않음을 의미하며, RQ2 의 distribution-aware 와 RQ3 의 distribution-agnostic 두 트랙을 모두 정당화하는 finding 이다.", {}),
])

add_p(doc, [
    ("이외에도 (v) q-error inf 발산, (vi) sample_size NaN 발산, (vii) Adaptive update path 의 SIGSEGV 가 hook 강제 활성화 직후 update 경로에서 발견되었다. 본 연구의 측정 단계에서는 ", {}),
    ("vector.update_sample_size = off", {"mono": True}),
    (" GUC 로 update 경로를 차단하고 sampling method 자체의 효과만 분리 측정하는 방식으로 회피하였다.", {}),
])

# RQ2 설정 [수정 7 — 오타]
add_h(doc, "RQ2 설정 및 검증 — 두 sanitize 의 직교 ablation", level=2)

add_p(doc, [
    ("RQ1 에서 드러난 단일 테이블 사각지대 위에서, 본 RQ 는 데이터 분포를 알 수 있을 때 두 가지 직교 sanitize 가 카디널리티 추정 정확도를 얼마나 개선하는지 측정한다.", {}),
])

add_p(doc, [
    ("RQ2-1 (Block sampling bias 제거).", {"bold": True}),
    (" ", {}),
    ("vector.c", {"mono": True}),
    (" 의 ", {}),
    ("TABLESAMPLE SYSTEM(p%)", {"mono": True}),
    (" 호출을 ", {}),
    ("TABLESAMPLE BERNOULLI(p%)", {"mono": True}),
    (" 로 한 줄 교체하여 행 단위 균일성을 회복한다. BERNOULLI 는 행 별 독립 베르누이 시행으로 추출하므로 block bias 가 제거된다.", {}),
])

add_p(doc, [
    ("RQ2-2 (Stratified sampling).", {"bold": True}),
    (" data-side 에서 사전에 k-means(K=20)으로 데이터를 20 개의 cluster 로 분할한 뒤, 기존 코드에 stratified sampling 분기를 추가한다. flag 활성화 시 각 cluster 별로 ", {}),
    ("ORDER BY random() LIMIT (sample_size / K)", {"mono": True}),
    (" 형태로 균등 샘플을 추출하고, Horvitz-Thompson 가중치(cluster 별 inclusion probability 의 역수)를 적용해 카디널리티를 계산한다. 본 절차는 RQ2-1 과 직교적이므로 ", {}),
    ("BERN/STRAT × dataset × selectivity", {"mono": True}),
    (" 의 직교 ablation matrix 로 두 sanitize 의 효과를 분리 측정할 수 있다.", {}),
])

# ─────────────────────────────────────────────────
# RQ2-1 검증 + 표 1 [수정 6]
# ─────────────────────────────────────────────────
add_h(doc, "RQ2-1 검증 — Block bias 제거 효과", level=3)

add_p(doc, [
    ("DEEP 1M 데이터셋에서 100 개 쿼리 × 6 개 selectivity 구간에 대해 SYSTEM ↔ BERNOULLI 만을 변화시킨 paired Wilcoxon 검정을 수행하였다. 표 1 은 EXPLAIN ANALYZE plan tree 의 ", {}),
    ("plan_rows", {"mono": True}),
    (" 를 reference 로 한 Python counterfactual 결과이며, native ", {}),
    ("vector.c", {"mono": True}),
    (" 한 줄 교체(Phase 4)로 측정한 같은 비교에서도 +3.8 ~ +9.6 % 의 같은 order of magnitude 개선이 확인되어 두 측정 방향이 일치한다.", {}),
])

# 표 1
add_p(doc, [("[표 1] Python counterfactual — Selectivity 별 SYSTEM vs BERNOULLI median Q-error 개선율 (DEEP 1M, 100 query, plan_rows 기반)", {"bold": True, "size": 10})])

t1 = doc.add_table(rows=7, cols=6)
t1.style = "Light Grid Accent 1"
hdrs = ["Selectivity", "System Median", "Bernoulli Median", "개선율", "p-value", "승리 (SYS<BERN)"]
for i, h in enumerate(hdrs):
    fill_cell(t1.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="D9E1F2")

rows = [
    ["0.001", "2.5970", "2.5970", "+0.0%",  "0.596",   "6/8"],
    ["0.010", "1.5584", "1.2987", "+20.0%", "0.240",   "45/40"],
    ["0.050", "1.2031", "1.1948", "+0.7%",  "0.0009",  "58/37"],
    ["0.100", "1.2120", "1.1323", "+7.0%",  "<0.001",  "66/31"],
    ["0.300", "1.2095", "1.0794", "+12.0%", "<0.001",  "76/24"],
    ["0.500", "1.1527", "1.0519", "+9.6%",  "<0.001",  "78/22"],
]
for ri, row in enumerate(rows, start=1):
    for ci, val in enumerate(row):
        fill_cell(t1.rows[ri].cells[ci], val, align=WD_ALIGN_PARAGRAPH.CENTER)

add_p(doc, [
    ("Selectivity 0.050 이상 4 개 구간에서 모두 paired Wilcoxon p < 0.001 (s = 0.050 만 0.0009)로 SYSTEM 의 Q-error 가 BERNOULLI 보다 통계적으로 유의하게 크다. 가장 큰 효과는 selectivity 0.500 에서 78/100 query 가 SYSTEM 쪽 불리, median 기준 약 9.6 %p 개선으로 나타났다. 이는 ", {}),
    ("vector.c", {"mono": True}),
    (" 한 줄 교체만으로도 Exqutor 의 카디널리티 추정 정확도가 측정 가능한 수준으로 개선될 수 있음을 시사한다.", {}),
])

# ─────────────────────────────────────────────────
# RQ2-2 검증 + 표 2 + Two-Level + HHI/CV [수정 5]
# ─────────────────────────────────────────────────
add_h(doc, "RQ2-2 검증 — Stratified sampling 의 추가 효과와 인과 분해", level=3)

add_p(doc, [
    ("KM20 stratified sampling 의 BERN 대비 추가 효과를 5-seed(setseed 0.1 ~ 0.5) 반복 측정으로 검증하였다. 표 2 는 세 데이터셋(DEEP 1M, DEEP 8M, SIFT 1.5M)에 대한 selectivity 0.500 / 0.050 / 0.010 의 5-seed 평균 개선율과 95 % 신뢰구간이다.", {}),
])

add_p(doc, [("[표 2] 5-seed 95 % 신뢰구간 — KM20 vs BERN median Q-error 개선율 (native vector.c 측정)", {"bold": True, "size": 10})])

t2 = doc.add_table(rows=4, cols=4)
t2.style = "Light Grid Accent 1"
hdrs2 = ["데이터셋", "s = 0.500", "s = 0.050", "s = 0.010"]
for i, h in enumerate(hdrs2):
    fill_cell(t2.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="D9E1F2")

rows2 = [
    ["DEEP 1M",   "+1.64 % [+1.11, +2.18]", "+1.85 % (단일)",          "+8.93 % (단일)"],
    ["DEEP 8M",   "+1.76 % [+0.65, +2.86]", "+0.55 % [-3.11, +4.21]",  "-0.71 % (노이즈)"],
    ["SIFT 1.5M", "+3.07 % [+2.66, +3.48]", "+4.39 % [+2.63, +6.15]",  "-0.53 % (양자화 영역)"],
]
for ri, row in enumerate(rows2, start=1):
    for ci, val in enumerate(row):
        fill_cell(t2.rows[ri].cells[ci], val, align=WD_ALIGN_PARAGRAPH.CENTER)

add_p(doc, [
    ("DEEP 1M 의 s = 0.500 / 0.300 / 0.100 세 구간이 모두 95 % 신뢰구간 하한이 양수로 확정되었으며([+1.11, +1.45, +0.72]), DEEP 8M 의 s = 0.500 신뢰구간 [+0.65, +2.86] 이 1M 의 [+1.11, +2.18] 과 겹쳐 데이터 규모에 대한 불변성이 확인되었다. SIFT 1.5M 의 s = 0.500 신뢰구간 [+2.66, +3.48] 은 DEEP 의 신뢰구간과 겹치지 않아, 더 쏠린 분포에서는 KM20 의 효과가 약 2 배로 정량 확대됨을 시사한다.", {}),
])

add_p(doc, [
    ("Two-Level Decomposition.", {"bold": True}),
    (" 본 팀은 KM20 의 개선분을 두 수준으로 분해하였다. ", {}),
    ("Level 1 (표본 안정화)", {"bold": True}),
    (" 은 단순히 데이터를 K = 20 으로 분할하여 비례 배분만 적용해도 발생하는 표본 크기 안정화 효과이며, ", {}),
    ("Level 2 (공간 인식)", {"bold": True}),
    (" 는 cluster 가 거리 기반(KM20)일 때만 추가로 발생하는 공간 인식 효과이다. DEEP 1M selectivity 0.010 에서 KM20 vs RANDOM20 의 격차는 19.6 %p 로, 좁은 selectivity 영역에서 무작위 파티션이 오히려 -10.67 % 로 악화되는 반면 KM20 은 +8.93 % 를 유지하여 Level 2 가 단독으로 +19.60 %p 를 기여함이 확인되었다. 이는 본 연구의 핵심 가설 — ", {}),
    ("데이터의 쏠림이 클수록 공간 인식 샘플링의 가치가 크다", {"italic": True}),
    (" — 를 직접 인과적으로 입증하는 anchor 결과이다.", {}),
])

add_p(doc, [
    ("HHI · CV 정량화.", {"bold": True}),
    (" SIFT 의 KM20 효과가 DEEP 의 약 2 배에 달하는 이유를 정량적으로 설명하기 위해, 본 팀은 각 데이터셋의 K = 20 cluster 크기 분포를 HHI(Herfindahl-Hirschman Index)와 CV(변동계수)로 측정하였다. DEEP 1M 의 cluster 크기 CV 는 0.234 인 반면 SIFT 1.5M 은 0.394 로 약 68 % 더 쏠려 있으며, 이러한 쏠림 격차가 KM20 효과의 약 2 배 격차로 단조 연결됨을 확인하였다. 이는 KM20 의 효과가 데이터셋 고유의 cluster 쏠림 분포로 사전 예측 가능함을 의미하며, RQ3(distribution-agnostic) 단계에서의 비교 baseline 으로도 활용된다.", {}),
])

# ─────────────────────────────────────────────────
# 마무리 — RQ3 설계 안내
# ─────────────────────────────────────────────────
add_h(doc, "RQ3 설계 — 본 중간보고서 시점의 설계 보고", level=2)

add_p(doc, [
    ("RQ3 는 데이터 분포를 사전에 모르는 상황에서, 사전 학습 없이 RQ2 의 공간 인식 효과를 얼마나 회수할 수 있는지를 ", {}),
    ("Recovery Rate", {"bold": True}),
    (" = (방법X − RANDOM20) / (KM20 − RANDOM20) 의 정의 위에서 측정한다. 본 중간보고서 시점에서는 다음 세 패러다임의 7 가지 방법이 설계만 확정된 상태이며, 실험은 최종보고서 단계(W5 ~ W7, 4/28 ~ 5/18)에서 수행한다.", {}),
])

add_p(doc, [
    ("Offline Partition.", {"bold": True}),
    (" 데이터 로드 시 1 회 계산하는 경량 사전 처리. A(LSH 랜덤 하이퍼플레인), C(Random Projection · Johnson-Lindenstrauss lemma), E(Hilbert Curve · 공간 DB 의 space-filling 기법을 벡터 카디널리티 추정에 최초 적용), F(Mini-batch K-means · 1 ~ 5 % 데이터만으로 근사 학습)의 4 종.", {}),
], style="List Bullet")

add_p(doc, [
    ("Online Query-Adaptive.", {"bold": True}),
    (" 쿼리 시점에 pilot sample 로 동적 적응. G(Distance-Shell · 쿼리 벡터 중심 동심원), B(KDE-pilot · Neyman 최적 배분)의 2 종.", {}),
], style="List Bullet")

add_p(doc, [
    ("Weight-based.", {"bold": True}),
    (" 파티션 없이 가중치로 분산을 감소. H(Importance Sampling)의 1 종.", {}),
], style="List Bullet")

# ─────────────────────────────────────────────────
# 출력
# ─────────────────────────────────────────────────
out = "/Users/hyunbin/Capstone/인종설 중간보고서_v2_조현빈반영.docx"
doc.save(out)
print(f"✅ saved: {out}")
