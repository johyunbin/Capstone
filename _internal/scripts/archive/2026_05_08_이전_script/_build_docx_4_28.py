"""
인종설 중간보고서 v3 minimalist (4/28 12:22 통합본)

피드백 반영:
- 4/27회의 (조현빈 정리): 분량 축소 + high-level + low-level 디테일 제거 + 한국어 위주 + RQ1 1/4 + 1만 배 이상
- 4/28회의 (이동욱 추가): 참고문헌에 Exqutor 추가, §4 제목 "연구 방법 및 결과", §4 4문단 맨앞 겹침 정리
- 4/27 raw data 검증: 표 2의 (단일 seed) → 5-seed t-based CI 정정

base: 이동욱 v2 (인종설 보고서-1.docx + 중간보고서_실험파트수정.docx) 학생 톤 그대로
제거: §3.2 코드 변경 위치 + 그림 1, §4.2.2 t-CI 산출식, §4.2.3 정량 (19.6%p / HHI), §4.3회복률·분모붕괴, 그림 6/7
유지/추가: 표 2의 (단일 seed) → CI 정정, "1만 배 이상", VAQ Analytical, 임채림 석사
신규: §5 진행상황 (학생 톤, 주 단위), §5.2 향후 일정 표, §6 역할 분담 표, 참고문헌
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

KOR = "Apple SD Gothic Neo"
ENG = "Helvetica Neue"
MONO = "Menlo"
FIG_DIR = "/Users/hyunbin/Capstone/experiments/figures"


def set_run_font(run, kor=KOR, eng=ENG, size=11, bold=False, italic=False):
    run.font.name = eng
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), kor)
    rFonts.set(qn("w:ascii"), eng)
    rFonts.set(qn("w:hAnsi"), eng)


def add_run(p, text, **kw):
    r = p.add_run(text)
    set_run_font(r, **kw)
    return r


def add_p(doc, text="", size=11, bold=False, align=None, space_before=4, space_after=4, line_spacing=1.0, indent_first=False, page_break_before=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if page_break_before:
        pf.page_break_before = True
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if indent_first:
        pf.first_line_indent = Cm(0.34)
    if align is not None:
        p.alignment = align
    if text:
        add_run(p, text, size=size, bold=bold)
    return p


def add_mixed_p(doc, runs, align=None, space_before=4, space_after=4, line_spacing=1.0, indent_first=True, page_break_before=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if page_break_before:
        pf.page_break_before = True
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if align is not None:
        p.alignment = align
    if indent_first:
        pf.first_line_indent = Cm(0.34)
    for text, kw in runs:
        add_run(p, text, **kw)
    return p


def add_heading_custom(doc, text, level=1, page_break_before=False):
    sizes = {1: 17, 2: 13, 3: 11.5}
    size = sizes.get(level, 11)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if page_break_before:
        pf.page_break_before = True
        pf.space_before = Pt(0)
    else:
        pf.space_before = Pt(8 if level == 1 else 10)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    add_run(p, text, size=size, bold=True)
    return p


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def fill_cell(cell, text, bold=False, align=None, mono=False, fill=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    eng = MONO if mono else ENG
    kor = MONO if mono else KOR
    add_run(p, text, kor=kor, eng=eng, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shade_cell(cell, fill)


def set_table_widths(t, widths_cm):
    for col, w in zip(t.columns, widths_cm):
        col.width = Cm(w)
    for row in t.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)


def fit_table_to_window(t):
    t.autofit = True
    tbl = t._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for tag in ("w:tblW", "w:tblLayout"):
        existing = tblPr.find(qn(tag))
        if existing is not None:
            tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), "5000")
    tblPr.append(tblW)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "autofit")
    tblPr.append(tblLayout)


def add_figure(doc, path, width_inches=6.27, caption=None):
    if not os.path.exists(path):
        add_p(doc, f"[그림 누락: {path}]", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(12)
    pf.keep_together = True
    img_run = p.add_run()
    img_run.add_picture(path, width=Inches(width_inches))
    if caption:
        spacer_run = p.add_run()
        spacer_run.add_break()
        spacer_run.add_break()
        for i, part in enumerate(caption.split("\n")):
            if i > 0:
                br = p.add_run()
                br.add_break()
            cap_run = p.add_run(part)
            set_run_font(cap_run, size=9.5)


def add_page_numbers(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_run_font(run, size=10)

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ─────────────────────────────────────────────────
doc = Document()


def enable_font_embedding(doc):
    settings = doc.settings.element
    for tag in ("embedTrueTypeFonts", "embedSystemFonts", "saveSubsetFonts"):
        existing = settings.find(qn(f"w:{tag}"))
        if existing is None:
            elem = OxmlElement(f"w:{tag}")
            elem.set(qn("w:val"), "true")
            settings.append(elem)


enable_font_embedding(doc)

style = doc.styles["Normal"]
style.font.name = ENG
style.font.size = Pt(11)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rPr.append(rFonts)
rFonts.set(qn("w:eastAsia"), KOR)
rFonts.set(qn("w:ascii"), ENG)
rFonts.set(qn("w:hAnsi"), ENG)

section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

# ═════════════════════════════════════════════════
# 표지
# ═════════════════════════════════════════════════

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "Skew-Aware Stratified Sampling을 이용한", size=22, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(22)
add_run(p, "벡터 카디널리티 추정 정확도 개선 연구", size=22, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
add_run(p, "인공지능종합설계 중간보고서", size=16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
add_run(p, "팀명 :  속도는벡터", size=14)

t = doc.add_table(rows=5, cols=2)
t.columns[0].width = Cm(3.5)
t.columns[1].width = Cm(6.0)
fill_cell(t.rows[0].cells[0], "팀     원", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="DDDDDD", size=11)
t.rows[0].cells[1].merge(t.rows[0].cells[0])
fill_cell(t.rows[0].cells[0], "팀     원", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="DDDDDD", size=11)
fill_cell(t.rows[1].cells[0], "팀장", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t.rows[1].cells[1], "박세은", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t.rows[2].cells[0], "팀원", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t.rows[2].cells[1], "강재현", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t.rows[3].cells[0], "팀원", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t.rows[3].cells[1], "이동욱", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t.rows[4].cells[0], "팀원", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t.rows[4].cells[1], "조현빈", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
for row in t.rows:
    row.height = Cm(1.0)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for cell in row.cells:
        for cp in cell.paragraphs:
            cp.paragraph_format.space_before = Pt(4)
            cp.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

t2 = doc.add_table(rows=2, cols=2)
t2.columns[0].width = Cm(3.5)
t2.columns[1].width = Cm(6.0)
fill_cell(t2.rows[0].cells[0], "지도교수", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t2.rows[0].cells[1], "박광현 교수님", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t2.rows[1].cells[0], "지도 연구원", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
fill_cell(t2.rows[1].cells[1], "임채림 석사", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
for row in t2.rows:
    row.height = Cm(1.0)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for cell in row.cells:
        for cp in cell.paragraphs:
            cp.paragraph_format.space_before = Pt(4)
            cp.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "2026 년 4 월 28일", size=12)

# ═════════════════════════════════════════════════
# Contents
# ═════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.page_break_before = True
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(36)
add_run(p, "Contents", size=28, bold=True)

contents = [
    ("1. 해결하고자 하는 문제", 1),
    ("2. 기존 연구의 현황 및 한계점", 1),
    ("3. 기존 연구와의 차별성 및 제안하는 연구의 중요성", 1),
    ("4. 연구 방법 및 결과", 1),
    ("     4.1 RQ1 — 사각지대의 구조적 한계 확인", 2),
    ("     4.2 RQ2 — 두 단계 보완의 직교 측정", 2),
    ("     4.3 RQ3 — 분포를 모르는 환경으로의 확장 (탐색 단계)", 2),
    ("5. 현재까지의 진행 상황 및 향후 계획", 1),
    ("6. 팀원별 역할 분담", 1),
    ("참고문헌", 1),
]
for i, (text, lv) in enumerate(contents):
    p = doc.add_paragraph()
    next_lv = contents[i + 1][1] if i + 1 < len(contents) else None
    if lv == 2 and next_lv == 1:
        p.paragraph_format.space_after = Pt(14)
    else:
        p.paragraph_format.space_after = Pt(11 if lv == 1 else 7)
    add_run(p, text, size=14 if lv == 1 else 12, bold=(lv == 1))

# ═════════════════════════════════════════════════
# 1. 해결하고자 하는 문제
# ═════════════════════════════════════════════════
add_heading_custom(doc, "1. 해결하고자 하는 문제", level=1, page_break_before=True)

add_heading_custom(doc, "1.1 벡터·임베딩과 분석 쿼리", level=2)

add_mixed_p(doc, [
    ("이미지·텍스트·음성과 같이 의미 기반의 비교가 필요한 데이터는 단순한 값 비교로 다루기 어렵다. 동일한 의미가 서로 다른 표현으로 나타나거나, 유사한 이미지라도 픽셀 값이 다르기 때문이다. 이를 해결하기 위해 최근에는 데이터를 임베딩 과정을 통해 의미 거리가 보존되는 고차원 벡터 (예: 96차원, 1024차원의 실수 배열)로 변환하는 방식이 널리 사용된다. 이때 의미적으로 유사한 항목일수록 벡터 공간에서도 가까운 위치에 놓이므로, 벡터 간 거리 계산을 통해 유사도를 효율적으로 추정할 수 있다.", {})
])

add_mixed_p(doc, [
    ("이와 같은 흐름 속에서 벡터 데이터베이스가 등장했으며, RAG (Retrieval-Augmented Generation), 멀티모달 검색, 추천 시스템과 같은 다양한 응용에서 핵심 인프라로 자리잡았다. 벡터 데이터베이스의 핵심 연산인 벡터 유사도 검색은 단순한 top-k 검색뿐 아니라, 거리 임계값 ", {}),
    ("D", {"eng": MONO, "kor": MONO}),
    ("이내의 모든 항목을 찾는 범위 검색 형태로도 자주 활용된다. 더 나아가 이러한 검색은 SQL의 조인·집계·필터 연산과 결합되어 벡터 증강 분석 쿼리 (VAQ, Vector-augmented Analytical Query)로 확장된다.", {})
])

add_heading_custom(doc, "1.2 카디널리티 추정과 옵티마이저", level=2)

add_mixed_p(doc, [
    ("위와 같은 쿼리를 효율적으로 실행하려면 옵티마이저가 각 연산의 결과 행 수, 즉 카디널리티를 정확히 추정해야 한다. 카디널리티 추정은 인덱스 사용 여부, 조인 알고리즘과 순서 등 실행 계획 전반에 직접적인 영향을 미친다. 잘못된 추정은 옵티마이저가 비효율적인 실행 계획을 선택하게 하고, 이는 곧 성능 저하로 이어진다. 따라서 정확한 카디널리티 추정은 데이터베이스 성능 최적화의 핵심 요소이다.", {})
])

# ═════════════════════════════════════════════════
# 2. 기존 연구의 현황 및 한계점
# ═════════════════════════════════════════════════
add_heading_custom(doc, "2. 기존 연구의 현황 및 한계점", level=1, page_break_before=True)

add_heading_custom(doc, "2.1 고정 비율 기반 선택도 추정", level=2)

add_mixed_p(doc, [
    ("pgvector[2]·VBASE[3]·DuckDB[4] 등 주요 벡터 데이터베이스 시스템은 거리 기반 조건의 선택도 (전체 행 중 조건을 만족하는 비율)를 각각 33.3%, 50%, 100%의 고정 비율로 추정한다. 그러나 실제 선택도는 쿼리에 따라 0.001%부터 100%까지 크게 변동한다. 고정 비율 추정은 데이터 분포와 쿼리 특성을 반영하지 못하며, 그 결과 옵티마이저가 비효율적인 실행 계획을 선택하는 문제가 발생한다. 실제로 잘못된 실행 계획은 수행 시간을 최대 1만 배 이상 증가시키기도 한다.", {})
])

add_heading_custom(doc, "2.2 Exqutor의 두 가지 보완책", level=2)

add_mixed_p(doc, [
    ("Exqutor[1]는 인덱스 존재 여부에 따라 두 가지 방식으로 이 문제를 보완한다.", {})
])

add_mixed_p(doc, [
    ("인덱스가 존재하는 경우 ", {}),
    ("ECQO (Exact Cardinality Query Optimization)", {"bold": True}),
    ("라는 방식을 사용한다. 실행 계획 수립 단계에서 실제 범위 검색을 소규모로 수행하여 정확한 카디널리티를 얻는다.", {})
])

add_mixed_p(doc, [
    ("인덱스가 존재하지 않는 경우 ", {}),
    ("Adaptive Sampling", {"bold": True}),
    (" (표본을 조금씩 늘려 가며 카디널리티를 근사하는 방식)이 일부 행만 추출해 임계값을 만족하는 비율을 계산하고, 그 비율을 전체 행 수에 곱해 카디널리티를 추산한다.", {})
])

add_mixed_p(doc, [
    ("Exqutor는 두 보완책을 통해 pgvector에서 최대 1,000배, VBASE에서 최대 1만 배 이상의 성능 향상을 보고한다.", {})
])

add_heading_custom(doc, "2.3 Exqutor의 한계점", level=2)

add_mixed_p(doc, [
    ("Exqutor는 유의미한 성능 개선을 보이지만 다음과 같은 한계를 가진다. 첫째, Exqutor의 보완책은 조인이 포함된 시나리오에서 작동하도록 구현되어 있어, 단일 테이블 검색에서는 Adaptive Sampling이 적용되지 않는다. 둘째, 샘플링이 특정 영역의 데이터에 편중될 수 있어 데이터 분포에 따라 추정 정확도가 크게 달라진다.", {})
])

# ═════════════════════════════════════════════════
# 3. 기존 연구와의 차별성 및 제안하는 연구의 중요성
# ═════════════════════════════════════════════════
add_heading_custom(doc, "3. 기존 연구와의 차별성 및 제안하는 연구의 중요성", level=1, page_break_before=True)

add_mixed_p(doc, [
    ("본 연구는 Exqutor가 가진 구조적 한계를 식별하고, 성능 저하의 원인을 분리하여 검증한다는 점에서 차별성 및 중요성을 가진다.", {})
])

add_mixed_p(doc, [
    ("첫째, 단일 테이블 시나리오에 집중한다. 실무에서는 하나의 테이블에서 조건을 만족하는 벡터를 찾는 단일 테이블 검색이 자주 등장한다. 그러나 Exqutor는 다중 테이블 조인을 전제로 설계되어, 이미지 검색이나 추천 시스템, RAG의 핵심인 단일 테이블 쿼리에서는 충분한 성능 개선을 보이지 못한다. 본 연구는 이러한 사각지대를 직접 다룬다.", {})
])

add_mixed_p(doc, [
    ("둘째, 기존 방법의 한계를 구분하여 분석한 후 각각의 요소에 맞는 해결책을 제시한다. Adaptive Sampling의 성능 저하는 크게 두 요인으로 설명될 수 있다. 하나는 연속된 물리적 저장 블록 단위로 표본을 가져오면서 특정 구간의 데이터가 과대표집되는 ", {}),
    ("block sampling bias", {}),
    ("이고, 다른 하나는 데이터의 공간적 불균형으로 인해 특정 영역이 샘플에서 누락되는 문제이다. 본 연구는 이를 완화하기 위해 데이터 공간을 여러 구간으로 나눈 뒤 각 구간에서 표본을 추출하는 ", {}),
    ("층화 표집 (stratified sampling)", {}),
    ("을 함께 검토한다. 두 요인은 서로 다른 단계에서 발생하므로, 실험적으로 분리하여 각각의 영향을 측정할 수 있다.", {})
])

# ═════════════════════════════════════════════════
# 4. 연구 방법 및 결과
# ═════════════════════════════════════════════════
add_heading_custom(doc, "4. 연구 방법 및 결과", level=1, page_break_before=True)

add_mixed_p(doc, [
    ("위의 문제를 바탕으로, 본 연구는 Adaptive Sampling의 구조적 한계를 분석하고, 개선 효과와 확장 가능성을 검증하기 위해 다음과 같은 연구 질문을 설정한다.", {})
])

add_mixed_p(doc, [
    ("RQ1. ", {"bold": True}),
    ("Adaptive Sampling은 단일 테이블 검색에서 왜 작동하지 않으며, 강제 적용 시 어떤 정확성 문제가 발생하는가?", {})
], indent_first=False)

add_mixed_p(doc, [
    ("RQ2. ", {"bold": True}),
    ("데이터 분포를 알고 있는 상황에서 block sampling bias 제거와 공간 인식 층화 표집은 카디널리티 추정 정확도를 얼마나 개선하는가?", {})
], indent_first=False)

add_mixed_p(doc, [
    ("RQ3. ", {"bold": True}),
    ("데이터 분포를 모르는 환경에서 공간 인식 표본화의 효과를 근사적으로 회수할 수 있는 후보 전략은 무엇인가?", {})
], indent_first=False)

add_heading_custom(doc, "4.1 RQ1 — 사각지대의 구조적 한계 확인", level=2)

add_mixed_p(doc, [
    ("RQ1은 Adaptive Sampling이 단일 테이블 검색에서 왜 작동하지 않으며, 강제 적용 시 어떤 정확성 문제가 발생하는지를 묻는다. Exqutor의 소스 코드를 분석한 결과, 기존 Adaptive Sampling은 주로 다중 테이블 조인 상황을 대상으로 설계되어 있어 단일 테이블 쿼리에서는 경로가 활성화되지 않는다는 점을 확인하였다. 이 경우 PostgreSQL의 기본 추정값이 사용되며, 벡터 조건의 실제 선택도를 제대로 반영하지 못한다.", {})
])

add_mixed_p(doc, [
    ("또한 단일 테이블 쿼리에서도 Adaptive Sampling이 실행되도록 조건을 수정하면, 표본 추출이 실제 쿼리 결과 자체에 영향을 주는 문제가 발생하였다. 예를 들어 실제로는 약 100,000개의 행이 반환되어야 하는 쿼리가 표본에 포함된 일부 행만을 대상으로 계산되어 32개의 결과를 반환하는 식이다. 이는 별도의 추정 경로가 필요함을 보여준다.", {})
])

add_mixed_p(doc, [
    ("추가로, 기존 Adaptive Sampling에서 사용하는 블록 단위 표본 추출 방식에도 한계가 있음을 확인하였다. 블록 단위 표본 추출은 데이터베이스 페이지 단위로 행을 함께 선택하거나 제외하기 때문에, 행 단위의 균일한 표본을 보장하지 못한다. 데이터가 특정 영역에 몰려 있는 경우 이러한 방식은 표본의 대표성을 낮추고 추정 오차를 키울 수 있다.", {})
])

add_mixed_p(doc, [
    ("마지막으로, 쿼리 정보만으로 데이터 분포의 쏠림 정도를 사전에 판단할 수 있는지도 검토하였다. 여러 글로벌·로컬 skewness 지표와 쿼리별 Q-error의 상관관계를 분석한 결과, 모든 조합에서 유의미한 상관관계가 관찰되지 않았다. 이는 쿼리 feature만으로 적절한 sampling 전략을 선택하기 어렵다는 것을 의미한다. 따라서 본 연구는 분포를 알고 있는 상황에서의 개선 방법 (RQ2)과, 분포를 모르는 상황으로의 확장 방법 (RQ3)을 구분하여 다룬다.", {})
])

# ═════════════════════════════════════════════════
# 4.2 RQ2
# ═════════════════════════════════════════════════
add_heading_custom(doc, "4.2 RQ2 — 두 단계 보완의 직교 측정", level=2, page_break_before=True)

add_mixed_p(doc, [
    ("RQ2에서는 데이터 분포를 알고 있는 상황을 가정하고, 두 가지 표본 추출 개선 방법의 효과를 측정하였다. 첫 번째는 기존 블록 단위 표본 추출을 행 단위 베르누이 표본 추출로 바꾸는 방법이다. 두 번째는 데이터를 공간적으로 비슷한 벡터끼리 묶은 뒤, 각 영역에서 표본을 추출하는 ", {}),
    ("k-means[5]", {}),
    (" 기반 층화 표본 방식이다.", {})
])

add_mixed_p(doc, [
    ("본 보고서에서는 기존 블록 단위 추출 방식을 SYSTEM, 행 단위 베르누이 추출 방식을 BERN, ", {}),
    ("k-means", {}),
    (" 기반 20개 클러스터 층화 표본을 KM20으로 표기한다. 평가 지표로는 실제 결과 행 수와 추정 결과 행 수의 차이를 비율로 나타내는 Q-error를 사용하였다. Q-error는 1에 가까울수록 추정 정확도가 높다.", {})
])

add_heading_custom(doc, "4.2.1 RQ2-1 — 블록 단위 표본 추출의 편향 완화", level=3)

add_mixed_p(doc, [
    ("RQ2-1에서는 SYSTEM과 BERN을 비교하여, 표본 추출 단위를 블록 단위에서 행 단위로 변경했을 때 추정 정확도가 개선되는지 확인하였다. 기존 SYSTEM 방식은 데이터베이스 페이지 단위로 표본을 추출하기 때문에 같은 블록 안의 행들이 함께 선택되거나 함께 제외된다. 이로 인해 행 단위의 균일한 표본 추출이 보장되지 않을 수 있다. 반면 BERN은 각 행을 독립적으로 표본에 포함할지 결정하므로, 행 단위 균일성을 더 잘 보장한다.", {})
])

add_mixed_p(doc, [
    ("실험은 DEEP 1M 데이터셋의 100개 query와 6개 selectivity 구간을 대상으로 수행하였다. 그 결과 selectivity 0.050 이상의 구간에서 BERN이 SYSTEM보다 전반적으로 낮은 Q-error를 보였다. 특히 selectivity 0.500에서는 median Q-error가 1.1527에서 1.0519로 감소하여 약 9.6%의 개선이 나타났고, 100개 query 중 78개에서 BERN이 더 좋은 결과를 보였다.", {})
])

add_mixed_p(doc, [
    ("반면 selectivity가 매우 낮은 구간에서는 두 방식의 차이가 뚜렷하지 않았다. 이 구간에서는 실제 결과 행 수가 작아 Q-error가 작은 변화에도 크게 변동될 수 있기 때문이다. 따라서 행 단위 추출은 중간 이상의 selectivity 구간에서 효과적이지만, 낮은 selectivity에서는 추가적인 보완이 필요하다.", {})
])

t1 = doc.add_table(rows=7, cols=6)
t1.style = "Light Grid Accent 1"
hdrs = ["Selectivity", "블록 단위", "행 단위", "개선율", "p-value", "행 단위 우위 (개수)"]
for i, h in enumerate(hdrs):
    fill_cell(t1.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="E7EBF2")
rows1 = [
    ["0.001", "2.5970", "2.5970", "+0.0%",  "0.596",   "6 / 8"],
    ["0.010", "1.5584", "1.2987", "+20.0%", "0.240",   "45 / 40"],
    ["0.050", "1.2031", "1.1948", "+0.7%",  "<0.001",  "58 / 37"],
    ["0.100", "1.2120", "1.1323", "+7.0%",  "<0.001",  "66 / 31"],
    ["0.300", "1.2095", "1.0794", "+12.0%", "<0.001",  "76 / 24"],
    ["0.500", "1.1527", "1.0519", "+9.6%",  "<0.001",  "78 / 22"],
]
for ri, row in enumerate(rows1, start=1):
    for ci, val in enumerate(row):
        fill_cell(t1.rows[ri].cells[ci], val, align=WD_ALIGN_PARAGRAPH.CENTER)
set_table_widths(t1, [2.4, 2.6, 2.6, 2.0, 2.2, 5.2])
fit_table_to_window(t1)

cap1 = add_p(doc, "표 1. Selectivity 별 추출 단위 변경 전후의 Q-error 중앙값 비교 (DEEP 1M, 100 쿼리)",
      bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
cap1.paragraph_format.space_before = Pt(14)
cap1_note = add_p(doc, "※ '행 단위 우위 (개수)'는 두 방식의 Q-error가 동일한 쿼리를 제외한 값이므로, 합이 100 미만일 수 있음.",
      bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

add_p(doc, "")

add_figure(doc, f"{FIG_DIR}/rq1_motivation/figure_1_phase4_scatter.png", width_inches=6.27,
           caption="그림 1. 추출 단위 변경 전후의 Q-error paired 비교 (DEEP 1M, selectivity 0.05·0.10·0.30·0.50)")

add_heading_custom(doc, "4.2.2 RQ2-2 — 공간 구조를 반영한 층화 표본", level=3, page_break_before=True)

add_mixed_p(doc, [
    ("RQ2-2에서는 BERN을 기준으로, 데이터의 공간 구조를 반영한 KM20 층화 표본이 추가적인 개선을 제공하는지 확인하였다. KM20은 전체 데이터를 ", {}),
    ("k-means", {}),
    ("로 20개 클러스터로 나눈 뒤, 각 클러스터에서 표본을 추출하고 클러스터 크기에 따라 가중치를 적용하여 전체 결과 행 수를 추정하는 방식이다.", {})
])

add_mixed_p(doc, [
    ("실험은 DEEP 1M, DEEP 8M, SIFT 1.5M 세 데이터셋의 selectivity 0.500 / 0.050 / 0.010 구간에서 5개 seed의 반복 측정으로 수행하였다. 그 결과 selectivity 0.500 구간에서는 세 데이터셋 모두에서 KM20이 BERN보다 추가적인 개선을 보였다. DEEP 1M에서는 +1.64%, DEEP 8M에서는 +1.76%, SIFT 1.5M에서는 +3.07%의 개선이 관찰되었으며, 95% 신뢰구간이 모두 양수로 확정되었다.", {})
])

add_mixed_p(doc, [
    ("특히 SIFT 1.5M에서 개선 효과가 더 크게 나타났는데, 이는 SIFT 데이터셋의 벡터 분포가 DEEP보다 더 불균등하기 때문으로 해석된다. 실제로 클러스터 크기 분포를 비교한 결과, SIFT의 클러스터 크기 변동계수는 DEEP보다 약 68% 더 컸다. 즉 데이터가 특정 영역에 더 강하게 몰려 있을수록, 공간 구조를 반영한 층화 표본의 효과가 커질 수 있음을 확인하였다.", {})
])

t2 = doc.add_table(rows=4, cols=4)
t2.style = "Light Grid Accent 1"
hdrs2 = ["데이터셋", "s = 0.500", "s = 0.050", "s = 0.010"]
for i, h in enumerate(hdrs2):
    fill_cell(t2.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="E7EBF2")
rows2 = [
    ["DEEP 1M",   "+1.64%  [+1.11, +2.18]", "+1.85%  [-0.22, +3.42]",     "+8.93%  [+6.59, +10.95]"],
    ["DEEP 8M",   "+1.76%  [+0.65, +2.86]", "+0.55%  [-3.11, +4.21]",     "-0.71%  (노이즈 영역)"],
    ["SIFT 1.5M", "+3.07%  [+2.66, +3.48]", "+4.39%  [+2.63, +6.15]",     "-0.53%  (양자화 영역)"],
]
for ri, row in enumerate(rows2, start=1):
    for ci, val in enumerate(row):
        fill_cell(t2.rows[ri].cells[ci], val, align=WD_ALIGN_PARAGRAPH.CENTER)
set_table_widths(t2, [2.6, 4.8, 4.8, 4.8])
fit_table_to_window(t2)

cap2 = add_p(doc, "표 2. 층화 표본의 RQ2-1 대비 추가 개선율 — 5-seed 평균과 95% 신뢰구간 (s = selectivity)",
      bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
cap2.paragraph_format.space_before = Pt(14)
cap2_note = add_p(doc, "※ 노이즈 영역: 결과 행 수가 작아 Q-error 분산이 신호를 압도. 양자화 영역: hook 추정값이 5-seed 모두 동일해 seed 간 분산이 사라진 영역.",
      bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

add_p(doc, "")

add_mixed_p(doc, [
    ("다만 selectivity 0.010의 일부 셀은 결과 행 수가 작아 Q-error가 불안정하게 변할 수 있으므로, 최종 보고서에서는 더 많은 반복 실험을 통해 해석을 보완할 필요가 있다.", {})
])

add_figure(doc, f"{FIG_DIR}/rq1_motivation/figure_2_phase6_box.png", width_inches=6.27,
           caption="그림 2. 추출 단위 변경 (BERN) 대비 층화 표본 (KM20)의 Q-error 분포 비교 (DEEP 1M, selectivity 6 구간)")

add_figure(doc, f"{FIG_DIR}/rq2_aware/figure_8_cross_dataset_bar.png", width_inches=6.27,
           caption="그림 3. 세 데이터셋의 selectivity 0.500 층화 표본 효과 (5-seed 95% 신뢰구간)")

add_figure(doc, f"{FIG_DIR}/rq2_aware/figure_7_selectivity_gradient.png", width_inches=6.27,
           caption="그림 4. 세 데이터셋의 selectivity 별 층화 표본 개선율 (selectivity 좌측 0.5 → 우측 0.01 로그 스케일)")

add_heading_custom(doc, "4.2.3 층화 표본 효과 분석", level=3, page_break_before=True)

add_mixed_p(doc, [
    ("추가 분석에서는 KM20의 개선 효과가 단순히 데이터를 여러 그룹으로 나누어서 발생하는 것인지, 아니면 ", {}),
    ("k-means", {}),
    ("가 벡터 공간의 구조를 반영했기 때문인지 확인하였다. 이를 위해 무작위로 20개 그룹을 나누는 RANDOM20과 비교하였다.", {})
])

add_mixed_p(doc, [
    ("DEEP 1M의 낮은 selectivity 조건에서 RANDOM20은 오히려 성능이 악화된 반면, KM20은 개선 효과를 유지하였다. 이는 벡터 공간의 구조를 반영한 분할이 중요하다는 점을 보여준다. 즉 KM20의 효과는 표본을 여러 그룹으로 나누는 안정화 효과뿐 아니라, 데이터의 공간적 쏠림을 반영하는 데서 추가로 발생한다.", {})
])

add_mixed_p(doc, [
    ("이 결과는 §4.2.2에서 관찰된 데이터셋별 효과 차이 (DEEP 대비 SIFT의 KM20 개선폭이 약 2배)와도 일관된다. 그림 6에 정리된 클러스터 크기 분포 비교에서 보이듯, 데이터의 공간적 쏠림 정도는 층화 표본 효과를 설명하는 핵심 요인으로 작용한다.", {})
])

add_figure(doc, f"{FIG_DIR}/rq2_aware/figure_9_two_level_decomposition.png", width_inches=6.27,
           caption="그림 5. 층화 표본의 Two-Level 분해 — 표본 안정화 (L1)와 공간 인식 (L2)의 분리")

add_figure(doc, f"{FIG_DIR}/rq2_aware/figure_10_cluster_skew.png", width_inches=6.27,
           caption="그림 6. DEEP 1M과 SIFT 1.5M의 20개 클러스터 크기 쏠림 비교")

# ═════════════════════════════════════════════════
# 4.3 RQ3
# ═════════════════════════════════════════════════
add_heading_custom(doc, "4.3 RQ3 — 분포를 모르는 환경으로의 확장 (탐색 단계)", level=2, page_break_before=True)

add_mixed_p(doc, [
    ("RQ2의 KM20 방식은 데이터 분포를 사전에 알고 있거나, 적어도 사전 클러스터링을 수행할 수 있다는 가정에 기반한다. 그러나 실제 운영 환경에서는 데이터가 계속 갱신되거나 분포를 미리 알기 어려울 수 있다. 따라서 RQ3는 분포를 모르는 상황에서도 KM20의 효과를 근사적으로 회수할 수 있는 방법을 탐구한다.", {})
])

add_mixed_p(doc, [
    ("후속 실험에서는 세 가지 유형의 방법을 비교할 예정이다. 첫째, 쿼리 실행 전에 데이터를 저비용으로 나누는 ", {}),
    ("offline partition", {}),
    (" 방식이다. 둘째, 쿼리가 들어온 시점에 동적으로 표본 배분을 조정하는 ", {}),
    ("online query-adaptive", {}),
    (" 방식이다. 셋째, 명시적으로 데이터를 나누지 않고 표본에 가중치를 부여하는 ", {}),
    ("importance sampling", {}),
    (" 방식이다. 후속 실험에서는 DEEP 1M, DEEP 8M, SIFT 1.5M 데이터셋을 대상으로 각 방법을 동일한 조건에서 비교할 계획이다.", {})
])

# ═════════════════════════════════════════════════
# 5. 현재까지의 진행 상황 및 향후 계획
# ═════════════════════════════════════════════════
add_heading_custom(doc, "5. 현재까지의 진행 상황 및 향후 계획", level=1, page_break_before=True)

add_heading_custom(doc, "5.1 진행 상황", level=2)

add_mixed_p(doc, [
    ("4월 첫째 주에는 지도 연구실로부터 Exqutor의 소스 코드와 실험 데이터셋, 실험 서버 권한을 인계받아 빌드와 환경 세팅을 마쳤다. 둘째 주에는 코드 분석을 통해 단일 테이블 사각지대를 확인하였고, 같은 시점에 출발 가설로 두었던 글로벌 쏠림 가설이 데이터에서 뒷받침되지 않는다는 점도 함께 확인하였다. 두 단계 보완 — 추출 단위를 행 단위로 바꾸는 단계와 공간 인식 층화 표본 단계 — 의 본 측정도 같은 주 안에 수행하였다.", {})
])

add_mixed_p(doc, [
    ("셋째 주에는 DEEP 8M과 SIFT 1.5M 두 데이터셋으로 외적 타당성을 확인하고, 5-seed 신뢰구간과 데이터셋별 쏠림 격차 측정을 마쳤다. 자문위원으로부터는 “단일 테이블 시나리오로 좁힌 점과 두 보완을 직교적으로 분리해서 잰 설계가 적절하다”는 회신을 받았다. 측정 과정에서 발견된 도구 오류는 자체 검증으로 즉시 식별·수정하였으며, 본 보고서의 모든 외적 타당성 수치는 재측정 결과이다.", {})
])

add_heading_custom(doc, "5.2 향후 계획", level=2)

t3 = doc.add_table(rows=4, cols=2)
t3.style = "Light Grid Accent 1"
for i, h in enumerate(["기간", "핵심 작업"]):
    fill_cell(t3.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="E7EBF2")
sched = [
    ["5/4 ~ 5/17",  "RQ3 보충 실험 및 통계 보강 (반복 측정, 신뢰구간 강화)"],
    ["5/18 ~ 5/24", "RQ3 분포 무인지 방법 비교 실험 (offline partition / online adaptive / importance sampling)"],
    ["5/25 ~ 6/14", "결과 분석, 최종 발표 (5/27 ~ 5/29), 전시회 (6/5), 최종 보고서 작성·제출 (6/11)"],
]
for ri, row in enumerate(sched, start=1):
    for ci, val in enumerate(row):
        fill_cell(t3.rows[ri].cells[ci], val, align=WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT)
set_table_widths(t3, [3.6, 13.4])
fit_table_to_window(t3)

cap3 = add_p(doc, "표 3. 중간 발표 이후 향후 일정",
      bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
cap3.paragraph_format.space_before = Pt(14)

add_p(doc, "")

add_mixed_p(doc, [
    ("중간 발표 이후의 5월은 RQ3의 후보 방법들을 동일한 측정 파이프라인에서 비교하는 작업에 집중하며, 6월에는 최종 발표·전시회·최종 보고서 작성을 마무리한다.", {})
])

# ═════════════════════════════════════════════════
# 6. 팀원별 역할 분담
# ═════════════════════════════════════════════════
add_heading_custom(doc, "6. 팀원별 역할 분담", level=1, page_break_before=True)

t4 = doc.add_table(rows=5, cols=2)
t4.style = "Light Grid Accent 1"
for i, h in enumerate(["팀원", "주요 역할"]):
    fill_cell(t4.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill="E7EBF2")
roles = [
    ["박세은 (팀장)",       "전체 일정 관리, 자문 회신 정리, 4인 합의 주재"],
    ["강재현 (주 발표자)",   "중간 발표 슬라이드 제작, 발표"],
    ["조현빈 (실험·분석)",   "실험 구현, 코드 수정 작업, 보고서 작성"],
    ["이동욱 (문서화·검토)", "보고서 작성, 보고서 및 발표 슬라이드 검토, 수정"],
]
for ri, row in enumerate(roles, start=1):
    for ci, val in enumerate(row):
        fill_cell(t4.rows[ri].cells[ci], val, align=WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT)
set_table_widths(t4, [3.6, 13.4])
fit_table_to_window(t4)

cap4 = add_p(doc, "표 4. 팀원별 역할 분담",
      bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
cap4.paragraph_format.space_before = Pt(14)

add_p(doc, "")

# ═════════════════════════════════════════════════
# 참고문헌
# ═════════════════════════════════════════════════
add_heading_custom(doc, "참고문헌", level=1, page_break_before=False)

add_mixed_p(doc, [
    ("[1]  BDAI Lab Research. Exqutor: Extended Query Optimizer for Vector-augmented Analytical Queries. arXiv:2512.09695v2, 2025.", {})
], indent_first=False)

add_mixed_p(doc, [
    ("[2]  A. Kane. pgvector: Open-source vector similarity search for Postgres. https://github.com/pgvector/pgvector, 2021.", {})
], indent_first=False)

add_mixed_p(doc, [
    ("[3]  Q. Wang et al. VBASE: Unifying Online Vector Similarity Search and Relational Queries via Relaxed Monotonicity. In OSDI, 2023.", {})
], indent_first=False)

add_mixed_p(doc, [
    ("[4]  M. Raasveldt and H. Mühleisen. DuckDB: an Embeddable Analytical Database. In SIGMOD, 2019.", {})
], indent_first=False)

add_mixed_p(doc, [
    ("[5]  S. Lloyd. Least squares quantization in PCM. IEEE Transactions on Information Theory, 28(2):129–137, 1982.", {})
], indent_first=False)

# ═════════════════════════════════════════════════
add_page_numbers(doc)

out = "/Users/hyunbin/Capstone/submission/_drafts/속도는벡터_중간보고서_20260428_183500.docx"
doc.save(out)
print(f"saved: {out}")
