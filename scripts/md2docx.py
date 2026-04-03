#!/usr/bin/env python3
"""
md2docx.py — Markdown → DOCX 변환 스크립트
사용법: python3 scripts/md2docx.py plans/문서이름.md
출력:   plans/문서이름.docx (같은 디렉토리)
방식:   Markdown 파싱 → python-docx로 DOCX 직접 생성
폰트:   Apple SD Gothic Neo (본문), D2Coding (코드)
"""

import re
import sys
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FONT_BODY = "Apple SD Gothic Neo"
FONT_CODE = "D2Coding"


def set_run_font(run, font_name=FONT_BODY, size=None, bold=False, italic=False, color=None):
    """Run의 폰트를 설정."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_font(paragraph, font_name=FONT_BODY, size=11):
    """Paragraph의 기본 폰트를 설정."""
    for run in paragraph.runs:
        set_run_font(run, font_name, size)


def add_heading_styled(doc, text, level):
    """스타일이 적용된 제목 추가."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = FONT_BODY
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    return heading


def add_table_from_rows(doc, headers, rows):
    """테이블 추가."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_run_font(run, size=9.5, bold=True)
        # 헤더 배경
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # 데이터
    for r_idx, row in enumerate(rows):
        for c_idx in range(len(headers)):
            cell = table.rows[r_idx + 1].cells[c_idx]
            val = row[c_idx] if c_idx < len(row) else ""
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=9.5)

    return table


def parse_inline(text):
    """인라인 마크다운(**bold**, *italic*, `code`) 파싱 → (text, bold, italic, code) 리스트."""
    parts = []
    # 패턴: **bold**, *italic*, `code`
    pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))"
    for match in re.finditer(pattern, text):
        if match.group(2):  # **bold**
            parts.append((match.group(2), True, False, False))
        elif match.group(3):  # *italic*
            parts.append((match.group(3), False, True, False))
        elif match.group(4):  # `code`
            parts.append((match.group(4), False, False, True))
        elif match.group(5):  # plain
            parts.append((match.group(5), False, False, False))
    return parts


def add_rich_paragraph(doc, text, size=11, alignment=None):
    """인라인 마크다운을 해석하여 paragraph에 추가."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment

    parts = parse_inline(text)
    for txt, bold, italic, code in parts:
        run = p.add_run(txt)
        if code:
            set_run_font(run, FONT_CODE, size=9.5, color=(229, 62, 62))
        else:
            set_run_font(run, FONT_BODY, size=size, bold=bold, italic=italic)

    return p


def convert(md_path):
    md_path = Path(md_path)
    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()

    # 기본 스타일 설정
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

    # 페이지 마진
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # 빈 줄
        if not line:
            i += 1
            continue

        # page-break div
        if '<div class="page-break">' in line:
            doc.add_page_break()
            i += 1
            # </div> 건너뛰기
            if i < len(lines) and "</div>" in lines[i]:
                i += 1
            continue

        # HTML div/table 건너뛰기 (inline HTML)
        if line.strip().startswith("<") and not line.strip().startswith("<div class=\"page-break\">"):
            # HTML 블록 건너뛰기
            while i < len(lines) and lines[i].strip():
                i += 1
            i += 1
            continue

        # --- (수평선)
        if line.strip() == "---":
            i += 1
            continue

        # 제목
        heading_match = re.match(r"^(#{1,4})\s+(.+)", line)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2)
            add_heading_styled(doc, title, level=min(level, 4))
            i += 1
            continue

        # 코드 블록
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # closing ```

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("\n".join(code_lines))
            set_run_font(run, FONT_CODE, size=9)
            # 코드 블록 배경
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0"/>')
            p._element.get_or_add_pPr().append(shading)
            continue

        # 테이블
        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|[-:|]+\|", lines[i + 1]):
            # 헤더 파싱
            headers = [c.strip() for c in line.split("|")[1:-1]]
            i += 2  # 헤더 + 구분선

            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                row = [c.strip() for c in lines[i].split("|")[1:-1]]
                rows.append(row)
                i += 1

            add_table_from_rows(doc, headers, rows)
            doc.add_paragraph()  # 테이블 뒤 간격
            continue

        # Blockquote
        if line.startswith("> "):
            quote_text = line[2:]
            i += 1
            while i < len(lines) and lines[i].startswith("> "):
                quote_text += " " + lines[i][2:]
                i += 1

            p = add_rich_paragraph(doc, quote_text, size=10.5)
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # 배경
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
            p._element.get_or_add_pPr().append(shading)
            continue

        # 리스트 (- 또는 숫자.)
        list_match = re.match(r"^(\s*)([-*]|\d+\.)\s+(.+)", line)
        if list_match:
            indent_level = len(list_match.group(1)) // 2
            bullet = list_match.group(2)
            content = list_match.group(3)

            if bullet in ["-", "*"]:
                p = doc.add_paragraph(style="List Bullet")
            else:
                p = doc.add_paragraph(style="List Number")

            # 기존 텍스트 지우고 rich text로 대체
            p.clear()
            parts = parse_inline(content)
            for txt, bold, italic, code in parts:
                run = p.add_run(txt)
                if code:
                    set_run_font(run, FONT_CODE, size=9.5, color=(229, 62, 62))
                else:
                    set_run_font(run, FONT_BODY, size=11, bold=bold, italic=italic)

            if indent_level > 0:
                p.paragraph_format.left_indent = Cm(1.2 * indent_level)

            i += 1
            continue

        # 메타 정보 (**key**: value)
        meta_match = re.match(r"^\*\*(.+?)\*\*:\s*(.+)", line)
        if meta_match:
            p = doc.add_paragraph()
            run_key = p.add_run(meta_match.group(1) + ": ")
            set_run_font(run_key, size=10, bold=True, color=(113, 128, 150))
            run_val = p.add_run(meta_match.group(2))
            set_run_font(run_val, size=10, color=(113, 128, 150))
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            i += 1
            continue

        # 일반 텍스트
        add_rich_paragraph(doc, line)
        i += 1

    # 저장
    out_path = md_path.with_suffix(".docx")
    doc.save(str(out_path))
    print(f"✓ {out_path}")
    return str(out_path)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("사용법: python3 scripts/md2docx.py <파일.md>")
        sys.exit(1)

    md_file = sys.argv[1]
    if not os.path.exists(md_file):
        print(f"파일 없음: {md_file}")
        sys.exit(1)

    convert(md_file)
